"""归并层仲裁设计范式 Demo。

设计范式说明：
1. 阶段 1：多个 chunk agent 并发分析。
   每个 agent 对自己负责的分块输出结构化意见，包括问题、修改建议、证据、关键条款值、待全局比对字段。
2. 阶段 2：主 agent 收集全部 chunk agent 的意见，形成全局意见池。
3. 阶段 3：主 agent 检测意见冲突。
   冲突类型包括：同一事项建议冲突、同一字段理解冲突、跨块联动冲突、漏判风险冲突。
4. 阶段 4：主 agent 仲裁。
   主 agent 决定保留、撤销、合并、联动标红的最终处理方案。
5. 阶段 5：输出最终报告。
   报告中的意见只保留仲裁后的结果，不直接输出原始 chunk agent 的意见列表。
"""

from __future__ import annotations

import asyncio
import json
import re
import tempfile
import time
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from types import SimpleNamespace
from typing import Any

from docx import Document

try:
    from .config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        ensure_multi_agent_demo_result_dir,
    )
    from .review_toolkit import MultiAgentReviewToolkit
except ImportError:
    from app.services.multi_agent.config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        ensure_multi_agent_demo_result_dir,
    )
    from app.services.multi_agent.review_toolkit import MultiAgentReviewToolkit


RISK_SCORE_MAP = {"高": 3, "中": 2, "低": 1}


@dataclass(slots=True)
class ChunkIssueOpinion:
    """单个 chunk agent 输出的问题意见。"""

    issue_id: str
    topic: str
    problem: str
    suggestion: str
    evidence: str
    risk_level: str
    key_clause_values: dict[str, str] = field(default_factory=dict)
    fields_to_compare: list[str] = field(default_factory=list)
    source_chunk_index: int = 0


@dataclass(slots=True)
class ChunkAnalysisOpinion:
    """单个 chunk agent 的完整结构化分析结果。"""

    chunk_index: int
    chunk_text: str
    chunk_summary: str
    issues: list[ChunkIssueOpinion] = field(default_factory=list)
    key_clause_values: dict[str, str] = field(default_factory=dict)
    fields_to_compare: list[str] = field(default_factory=list)
    raw_output: str = ""


@dataclass(slots=True)
class ConflictRecord:
    """全局冲突记录。"""

    conflict_id: str
    conflict_type: str
    title: str
    description: str
    chunk_indices: list[int] = field(default_factory=list)
    related_issue_ids: list[str] = field(default_factory=list)
    field_name: str = ""
    topic: str = ""
    values: list[str] = field(default_factory=list)


@dataclass(slots=True)
class ArbitrationDecision:
    """主 agent 仲裁决策。"""

    decision_id: str
    action: str
    title: str
    rationale: str
    kept_issue_ids: list[str] = field(default_factory=list)
    withdrawn_issue_ids: list[str] = field(default_factory=list)
    highlight_chunks: list[int] = field(default_factory=list)


@dataclass(slots=True)
class FinalArbitratedIssue:
    """仲裁后的最终意见。"""

    final_issue_id: str
    title: str
    risk_level: str
    problem: str
    suggestion: str
    evidence: str
    source_chunks: list[int] = field(default_factory=list)
    source_issue_ids: list[str] = field(default_factory=list)


class MergeArbitrationReviewDemo(MultiAgentReviewToolkit):
    """带归并层与仲裁层的多 agent 审阅 demo。"""

    async def run(
        self,
        file_path: str,
        *,
        stance: str | None = None,
        intensity: str | None = None,
        contract_type: str | None = None,
    ) -> dict[str, Any]:
        """执行归并层仲裁 demo。"""
        start_time = time.perf_counter()
        parsed_contract = await self.parse_contract_file(file_path)
        chunks = self.split_contract_text(parsed_contract.text)

        chunk_opinions = await self._run_stage_one_parallel_analysis(
            chunks=chunks,
            stance=stance or self.demo_config.default_stance,
            intensity=intensity or self.demo_config.default_intensity,
            contract_type=contract_type or self.demo_config.default_contract_type,
        )
        pooled_issues = self._collect_global_opinion_pool(chunk_opinions)
        conflicts = self._detect_conflicts(chunk_opinions, pooled_issues)
        decisions, final_issues = self._arbitrate_conflicts(
            chunk_opinions=chunk_opinions,
            pooled_issues=pooled_issues,
            conflicts=conflicts,
        )
        elapsed_seconds = time.perf_counter() - start_time
        result_file_path = self._save_final_report(
            parsed_contract=parsed_contract,
            chunk_count=len(chunks),
            pooled_issue_count=len(pooled_issues),
            conflicts=conflicts,
            decisions=decisions,
            final_issues=final_issues,
            elapsed_seconds=elapsed_seconds,
        )

        return {
            "demo_type": "merge_arbitration",
            "file_path": parsed_contract.file_path,
            "file_type": parsed_contract.file_type,
            "chunk_count": len(chunks),
            "pooled_issue_count": len(pooled_issues),
            "conflict_count": len(conflicts),
            "final_issue_count": len(final_issues),
            "elapsed_seconds": elapsed_seconds,
            "result_file_path": result_file_path,
        }

    async def _run_stage_one_parallel_analysis(
        self,
        *,
        chunks: list[str],
        stance: str,
        intensity: str,
        contract_type: str,
    ) -> list[ChunkAnalysisOpinion]:
        """阶段 1：并发执行 chunk agent 分析。"""
        semaphore = asyncio.Semaphore(self.demo_config.max_concurrent_reviews)
        ordered_results: list[ChunkAnalysisOpinion | None] = [None] * len(chunks)

        async def _analyze_single(index: int, chunk_text: str):
            async with semaphore:
                context = f"这是第 {index + 1} 个分块，共 {len(chunks)} 个。"
                ordered_results[index] = await self._analyze_chunk_with_agent(
                    chunk_index=index + 1,
                    chunk_text=chunk_text,
                    stance=stance,
                    intensity=intensity,
                    contract_type=contract_type,
                    context=context,
                )

        await asyncio.gather(
            *[_analyze_single(index, chunk) for index, chunk in enumerate(chunks)]
        )
        return [item for item in ordered_results if item is not None]

    async def _analyze_chunk_with_agent(
        self,
        *,
        chunk_index: int,
        chunk_text: str,
        stance: str,
        intensity: str,
        contract_type: str,
        context: str,
    ) -> ChunkAnalysisOpinion:
        """阶段 1：单个 chunk agent 的结构化分析。"""
        prompt = self._build_stage_one_prompt(
            chunk_text=chunk_text,
            stance=stance,
            intensity=intensity,
            contract_type=contract_type,
            context=context,
        )
        messages = [
            {"role": "system", "content": "你是合同审阅多 agent 系统中的分块分析 agent。"},
            {"role": "user", "content": prompt},
        ]
        raw_output = ""
        try:
            response = await asyncio.to_thread(
                self.llm.chat.completions.create,
                model=self.model_config.model_name,
                messages=messages,
                tools=[],
                temperature=self.model_config.temperature,
                top_p=self.model_config.top_p,
                max_tokens=self.model_config.max_tokens,
            )
            raw_output = response.choices[0].message.content or ""
            return self._parse_stage_one_output(
                chunk_index=chunk_index,
                chunk_text=chunk_text,
                raw_output=raw_output,
            )
        except Exception:
            return self._fallback_stage_one_output(
                chunk_index=chunk_index,
                chunk_text=chunk_text,
                raw_output=raw_output,
            )

    def _build_stage_one_prompt(
        self,
        *,
        chunk_text: str,
        stance: str,
        intensity: str,
        contract_type: str,
        context: str,
    ) -> str:
        """构造阶段 1 的结构化分析提示词。"""
        return (
            "你是合同审阅系统中的分块分析 agent，请只分析当前分块，并以纯 JSON 输出。\n"
            "请严格输出以下 JSON 结构，不要输出 markdown，不要输出解释文字：\n"
            "{\n"
            '  "chunk_summary": "字符串",\n'
            '  "issues": [\n'
            "    {\n"
            '      "topic": "问题主题",\n'
            '      "problem": "本块原文中的问题",\n'
            '      "suggestion": "修改建议",\n'
            '      "evidence": "支持该判断的原文证据",\n'
            '      "risk_level": "高/中/低",\n'
            '      "key_clause_values": {"字段名": "字段值"},\n'
            '      "fields_to_compare": ["需要和全局比对的字段"]\n'
            "    }\n"
            "  ],\n"
            '  "key_clause_values": {"字段名": "字段值"},\n'
            '  "fields_to_compare": ["需要和全局比对的字段"]\n'
            "}\n"
            "要求：\n"
            "1. 若本块无明显风险，issues 返回空数组，但仍需提取关键条款值与待比对字段。\n"
            "2. evidence 必须引用本块中的关键证据，不得编造。\n"
            "3. key_clause_values 只保留适合做跨块比对的关键字段，例如付款期限、验收条件、违约金比例、工期、生效条件、争议解决地。\n"
            "4. fields_to_compare 必须列出需要在全局范围核验一致性的字段。\n"
            f"5. 用户立场：{stance}；审阅强度：{intensity}；合同类型：{contract_type}。\n"
            f"6. 审阅上下文：{context}\n"
            "以下是当前分块内容：\n"
            f"{chunk_text}"
        )

    def _parse_stage_one_output(
        self,
        *,
        chunk_index: int,
        chunk_text: str,
        raw_output: str,
    ) -> ChunkAnalysisOpinion:
        """解析阶段 1 的 JSON 输出。"""
        payload = self._extract_json_payload(raw_output)
        chunk_summary = str(payload.get("chunk_summary", "")).strip()
        chunk_values = self._normalize_clause_values(payload.get("key_clause_values"))
        chunk_compare_fields = self._normalize_fields(payload.get("fields_to_compare"))

        issues: list[ChunkIssueOpinion] = []
        raw_issues = payload.get("issues") or []
        for issue_index, item in enumerate(raw_issues, start=1):
            topic = str(item.get("topic", "")).strip() or f"问题{issue_index}"
            problem = str(item.get("problem", "")).strip()
            suggestion = str(item.get("suggestion", "")).strip()
            evidence = str(item.get("evidence", "")).strip()
            risk_level = self._normalize_risk_level(item.get("risk_level"))
            issue_values = self._normalize_clause_values(item.get("key_clause_values"))
            issue_fields = self._normalize_fields(item.get("fields_to_compare"))

            issues.append(
                ChunkIssueOpinion(
                    issue_id=f"chunk-{chunk_index}-issue-{issue_index}",
                    topic=topic,
                    problem=problem,
                    suggestion=suggestion,
                    evidence=evidence,
                    risk_level=risk_level,
                    key_clause_values=issue_values,
                    fields_to_compare=issue_fields,
                    source_chunk_index=chunk_index,
                )
            )

        return ChunkAnalysisOpinion(
            chunk_index=chunk_index,
            chunk_text=chunk_text,
            chunk_summary=chunk_summary,
            issues=issues,
            key_clause_values=chunk_values,
            fields_to_compare=chunk_compare_fields,
            raw_output=raw_output,
        )

    def _fallback_stage_one_output(
        self,
        *,
        chunk_index: int,
        chunk_text: str,
        raw_output: str,
    ) -> ChunkAnalysisOpinion:
        """阶段 1 输出解析失败时的兜底。"""
        parsed_modifications = self._parse_review_result(raw_output) if raw_output else []
        issues: list[ChunkIssueOpinion] = []
        for issue_index, item in enumerate(parsed_modifications, start=1):
            issues.append(
                ChunkIssueOpinion(
                    issue_id=f"chunk-{chunk_index}-issue-{issue_index}",
                    topic=str(item.get("risk_type", "")).strip() or f"问题{issue_index}",
                    problem=str(item.get("risk_analysis", "")).strip(),
                    suggestion=str(item.get("suggested_content", "")).strip(),
                    evidence=str(item.get("original_content", "")).strip(),
                    risk_level=self._normalize_risk_level(item.get("risk_level")),
                    key_clause_values={},
                    fields_to_compare=[],
                    source_chunk_index=chunk_index,
                )
            )
        return ChunkAnalysisOpinion(
            chunk_index=chunk_index,
            chunk_text=chunk_text,
            chunk_summary="结构化解析失败，已退化为普通审阅结果。",
            issues=issues,
            key_clause_values={},
            fields_to_compare=[],
            raw_output=raw_output,
        )

    def _collect_global_opinion_pool(
        self,
        chunk_opinions: list[ChunkAnalysisOpinion],
    ) -> list[ChunkIssueOpinion]:
        """阶段 2：收集全部 chunk agent 的意见池。"""
        return [issue for opinion in chunk_opinions for issue in opinion.issues]

    def _detect_conflicts(
        self,
        chunk_opinions: list[ChunkAnalysisOpinion],
        pooled_issues: list[ChunkIssueOpinion],
    ) -> list[ConflictRecord]:
        """阶段 3：检测意见冲突。"""
        conflicts: list[ConflictRecord] = []
        conflict_index = 1

        field_conflicts, conflicted_fields = self._detect_field_value_conflicts(
            chunk_opinions,
            pooled_issues,
            start_index=conflict_index,
        )
        conflicts.extend(field_conflicts)
        conflict_index += len(field_conflicts)

        suggestion_conflicts = self._detect_suggestion_conflicts(
            pooled_issues,
            start_index=conflict_index,
        )
        conflicts.extend(suggestion_conflicts)
        conflict_index += len(suggestion_conflicts)

        cross_chunk_conflicts = self._detect_cross_chunk_conflicts(
            pooled_issues,
            conflicted_fields=conflicted_fields,
            start_index=conflict_index,
        )
        conflicts.extend(cross_chunk_conflicts)
        conflict_index += len(cross_chunk_conflicts)

        risk_contradictions = self._detect_risk_missed_conflicts(
            chunk_opinions,
            conflicted_fields=conflicted_fields,
            start_index=conflict_index,
        )
        conflicts.extend(risk_contradictions)
        return conflicts

    def _detect_field_value_conflicts(
        self,
        chunk_opinions: list[ChunkAnalysisOpinion],
        pooled_issues: list[ChunkIssueOpinion],
        *,
        start_index: int,
    ) -> tuple[list[ConflictRecord], set[str]]:
        """检测关键字段值冲突。"""
        field_entries: dict[str, list[tuple[int, str, str]]] = defaultdict(list)
        field_issue_ids: dict[str, list[str]] = defaultdict(list)

        for opinion in chunk_opinions:
            for field_name, field_value in opinion.key_clause_values.items():
                if field_value:
                    normalized = self._normalize_key(field_name)
                    field_entries[normalized].append((opinion.chunk_index, field_name, field_value))

        for issue in pooled_issues:
            for field_name, field_value in issue.key_clause_values.items():
                if field_value:
                    normalized = self._normalize_key(field_name)
                    field_entries[normalized].append((issue.source_chunk_index, field_name, field_value))
                    field_issue_ids[normalized].append(issue.issue_id)

        conflicts: list[ConflictRecord] = []
        conflicted_fields: set[str] = set()
        current_index = start_index
        for normalized_field, entries in field_entries.items():
            distinct_values = {
                self._normalize_value(item[2]): item[2]
                for item in entries
                if self._normalize_value(item[2])
            }
            if len(distinct_values) <= 1:
                continue
            conflicted_fields.add(normalized_field)
            display_field = next((item[1] for item in entries if item[1].strip()), normalized_field)
            conflicts.append(
                ConflictRecord(
                    conflict_id=f"conflict-{current_index}",
                    conflict_type="field_value_inconsistency",
                    title=f"关键字段冲突：{display_field}",
                    description=f"多个分块对字段“{display_field}”提取出了不一致的条款值，需要全局仲裁。",
                    chunk_indices=sorted({item[0] for item in entries}),
                    related_issue_ids=sorted(set(field_issue_ids.get(normalized_field, []))),
                    field_name=display_field,
                    values=list(distinct_values.values()),
                )
            )
            current_index += 1
        return conflicts, conflicted_fields

    def _detect_suggestion_conflicts(
        self,
        pooled_issues: list[ChunkIssueOpinion],
        *,
        start_index: int,
    ) -> list[ConflictRecord]:
        """检测同一事项建议不一致。"""
        topic_map: dict[str, list[ChunkIssueOpinion]] = defaultdict(list)
        for issue in pooled_issues:
            normalized_topic = self._normalize_key(issue.topic or issue.problem[:20])
            if normalized_topic:
                topic_map[normalized_topic].append(issue)

        conflicts: list[ConflictRecord] = []
        current_index = start_index
        for normalized_topic, issues in topic_map.items():
            distinct_suggestions = {
                self._normalize_value(issue.suggestion): issue.suggestion
                for issue in issues
                if self._normalize_value(issue.suggestion)
            }
            if len(distinct_suggestions) <= 1:
                continue
            display_topic = next((issue.topic for issue in issues if issue.topic.strip()), normalized_topic)
            conflicts.append(
                ConflictRecord(
                    conflict_id=f"conflict-{current_index}",
                    conflict_type="suggestion_inconsistency",
                    title=f"同一事项建议冲突：{display_topic}",
                    description=f"多个分块围绕“{display_topic}”提出了不同修改建议，需要主 agent 决定保留或合并方案。",
                    chunk_indices=sorted({issue.source_chunk_index for issue in issues}),
                    related_issue_ids=[issue.issue_id for issue in issues],
                    topic=display_topic,
                    values=list(distinct_suggestions.values()),
                )
            )
            current_index += 1
        return conflicts

    def _detect_cross_chunk_conflicts(
        self,
        pooled_issues: list[ChunkIssueOpinion],
        *,
        conflicted_fields: set[str],
        start_index: int,
    ) -> list[ConflictRecord]:
        """检测跨块联动冲突。"""
        grouped: dict[str, list[ChunkIssueOpinion]] = defaultdict(list)
        for issue in pooled_issues:
            overlap_fields = [
                field_name
                for field_name in issue.fields_to_compare
                if self._normalize_key(field_name) in conflicted_fields
            ]
            if not overlap_fields:
                continue
            key = "|".join(sorted({self._normalize_key(field) for field in overlap_fields}))
            grouped[key].append(issue)

        conflicts: list[ConflictRecord] = []
        current_index = start_index
        for key, issues in grouped.items():
            involved_fields = sorted(
                {
                    field_name
                    for issue in issues
                    for field_name in issue.fields_to_compare
                    if self._normalize_key(field_name) in conflicted_fields
                }
            )
            conflicts.append(
                ConflictRecord(
                    conflict_id=f"conflict-{current_index}",
                    conflict_type="cross_chunk_linked_change",
                    title=f"跨块联动冲突：{'、'.join(involved_fields)}",
                    description="部分修改建议会直接影响其他分块中的对应字段，需要形成跨块联动修改建议。",
                    chunk_indices=sorted({issue.source_chunk_index for issue in issues}),
                    related_issue_ids=[issue.issue_id for issue in issues],
                    values=involved_fields,
                )
            )
            current_index += 1
        return conflicts

    def _detect_risk_missed_conflicts(
        self,
        chunk_opinions: list[ChunkAnalysisOpinion],
        *,
        conflicted_fields: set[str],
        start_index: int,
    ) -> list[ConflictRecord]:
        """检测局部无风险但全局有风险的冲突。"""
        risky_fields = set(conflicted_fields)
        for opinion in chunk_opinions:
            risky_fields.update(self._normalize_key(field) for field in opinion.fields_to_compare)

        conflicts: list[ConflictRecord] = []
        current_index = start_index
        for opinion in chunk_opinions:
            if opinion.issues:
                continue
            local_fields = {
                self._normalize_key(field)
                for field in opinion.fields_to_compare + list(opinion.key_clause_values.keys())
            }
            overlap = sorted(field for field in local_fields if field in risky_fields)
            if not overlap:
                continue
            conflicts.append(
                ConflictRecord(
                    conflict_id=f"conflict-{current_index}",
                    conflict_type="risk_missed",
                    title=f"漏判风险冲突：分块{opinion.chunk_index}",
                    description="该分块自身未识别出明确风险，但它涉及的关键字段在其他分块出现了冲突或显著风险，需联动复核。",
                    chunk_indices=[opinion.chunk_index],
                    related_issue_ids=[],
                    values=overlap,
                )
            )
            current_index += 1
        return conflicts

    def _arbitrate_conflicts(
        self,
        *,
        chunk_opinions: list[ChunkAnalysisOpinion],
        pooled_issues: list[ChunkIssueOpinion],
        conflicts: list[ConflictRecord],
    ) -> tuple[list[ArbitrationDecision], list[FinalArbitratedIssue]]:
        """阶段 4：主 agent 仲裁冲突。"""
        issue_map = {issue.issue_id: issue for issue in pooled_issues}
        withdrawn_issue_ids: set[str] = set()
        covered_issue_ids: set[str] = set()
        decisions: list[ArbitrationDecision] = []
        final_issues: list[FinalArbitratedIssue] = []
        final_issue_index = 1

        for conflict in conflicts:
            related_issues = [
                issue_map[issue_id]
                for issue_id in conflict.related_issue_ids
                if issue_id in issue_map and issue_id not in withdrawn_issue_ids
            ]

            if conflict.conflict_type == "field_value_inconsistency":
                retained_ids = [issue.issue_id for issue in related_issues]
                withdrawn_issue_ids.update(retained_ids)
                covered_issue_ids.update(retained_ids)
                evidence = "；".join(
                    f"分块{issue.source_chunk_index}：{issue.evidence or issue.problem}"
                    for issue in related_issues
                ) or "多个分块提取的关键字段值不一致。"
                final_issues.append(
                    FinalArbitratedIssue(
                        final_issue_id=f"final-{final_issue_index}",
                        title=conflict.title,
                        risk_level="高",
                        problem=conflict.description,
                        suggestion=(
                            f"对字段“{conflict.field_name}”形成统一口径，并联动修改相关分块；"
                            f"当前检测到的冲突值包括：{'；'.join(conflict.values)}。"
                        ),
                        evidence=evidence,
                        source_chunks=conflict.chunk_indices,
                        source_issue_ids=retained_ids,
                    )
                )
                final_issue_index += 1
                decisions.append(
                    ArbitrationDecision(
                        decision_id=f"decision-{len(decisions) + 1}",
                        action="merge",
                        title=conflict.title,
                        rationale="该冲突涉及同一关键字段的全局一致性，应合并为跨块联动修改建议统一处理。",
                        kept_issue_ids=[],
                        withdrawn_issue_ids=retained_ids,
                        highlight_chunks=conflict.chunk_indices,
                    )
                )
                continue

            if conflict.conflict_type == "suggestion_inconsistency" and related_issues:
                primary_issue = max(
                    related_issues,
                    key=lambda item: (
                        RISK_SCORE_MAP.get(item.risk_level, 0),
                        len(item.evidence),
                    ),
                )
                alternative_issues = [
                    issue for issue in related_issues if issue.issue_id != primary_issue.issue_id
                ]
                withdrawn_ids = [issue.issue_id for issue in alternative_issues]
                withdrawn_issue_ids.update(withdrawn_ids)
                covered_issue_ids.update(conflict.related_issue_ids)
                merged_evidence = "；".join(
                    f"分块{issue.source_chunk_index}：{issue.evidence or issue.problem}"
                    for issue in related_issues
                )
                final_issues.append(
                    FinalArbitratedIssue(
                        final_issue_id=f"final-{final_issue_index}",
                        title=conflict.title,
                        risk_level=primary_issue.risk_level,
                        problem=conflict.description,
                        suggestion=primary_issue.suggestion,
                        evidence=merged_evidence,
                        source_chunks=sorted({issue.source_chunk_index for issue in related_issues}),
                        source_issue_ids=[issue.issue_id for issue in related_issues],
                    )
                )
                final_issue_index += 1
                decisions.append(
                    ArbitrationDecision(
                        decision_id=f"decision-{len(decisions) + 1}",
                        action="keep_one_and_merge_context",
                        title=conflict.title,
                        rationale="同一事项出现多个修改版本时，优先保留风险等级更高、证据更充分的建议，并吸收其他分块证据。",
                        kept_issue_ids=[primary_issue.issue_id],
                        withdrawn_issue_ids=withdrawn_ids,
                        highlight_chunks=sorted({issue.source_chunk_index for issue in related_issues}),
                    )
                )
                continue

            if conflict.conflict_type == "cross_chunk_linked_change":
                related_ids = [issue.issue_id for issue in related_issues]
                covered_issue_ids.update(related_ids)
                final_issues.append(
                    FinalArbitratedIssue(
                        final_issue_id=f"final-{final_issue_index}",
                        title=conflict.title,
                        risk_level="高",
                        problem=conflict.description,
                        suggestion=(
                            f"将涉及字段“{'、'.join(conflict.values)}”的修改建议统一编排，"
                            "避免一个分块改动后与其他分块形成新的条款冲突。"
                        ),
                        evidence="；".join(
                            f"分块{issue.source_chunk_index}：{issue.problem}"
                            for issue in related_issues
                        ) or conflict.description,
                        source_chunks=conflict.chunk_indices,
                        source_issue_ids=related_ids,
                    )
                )
                final_issue_index += 1
                decisions.append(
                    ArbitrationDecision(
                        decision_id=f"decision-{len(decisions) + 1}",
                        action="merge_as_cross_chunk_change",
                        title=conflict.title,
                        rationale="涉及全局比对字段的修改不能逐块独立处理，应上升为跨块联动修改项。",
                        kept_issue_ids=[],
                        withdrawn_issue_ids=[],
                        highlight_chunks=conflict.chunk_indices,
                    )
                )
                continue

            if conflict.conflict_type == "risk_missed":
                final_issues.append(
                    FinalArbitratedIssue(
                        final_issue_id=f"final-{final_issue_index}",
                        title=conflict.title,
                        risk_level="中",
                        problem=conflict.description,
                        suggestion="将该分块与相关高风险字段一并复核，并在最终报告中联动标注。",
                        evidence=f"该分块关联的全局风险字段：{'、'.join(conflict.values)}。",
                        source_chunks=conflict.chunk_indices,
                        source_issue_ids=[],
                    )
                )
                final_issue_index += 1
                decisions.append(
                    ArbitrationDecision(
                        decision_id=f"decision-{len(decisions) + 1}",
                        action="highlight_for_recheck",
                        title=conflict.title,
                        rationale="虽然单块未识别风险，但全局证据表明该块应联动复核。",
                        kept_issue_ids=[],
                        withdrawn_issue_ids=[],
                        highlight_chunks=conflict.chunk_indices,
                    )
                )

        for issue in pooled_issues:
            if issue.issue_id in withdrawn_issue_ids:
                continue
            if issue.issue_id in covered_issue_ids:
                continue
            final_issues.append(
                FinalArbitratedIssue(
                    final_issue_id=f"final-{final_issue_index}",
                    title=issue.topic,
                    risk_level=issue.risk_level,
                    problem=issue.problem,
                    suggestion=issue.suggestion,
                    evidence=issue.evidence,
                    source_chunks=[issue.source_chunk_index],
                    source_issue_ids=[issue.issue_id],
                )
            )
            final_issue_index += 1

        final_issues.sort(
            key=lambda item: (
                -RISK_SCORE_MAP.get(item.risk_level, 0),
                item.final_issue_id,
            )
        )
        return decisions, final_issues

    def _save_final_report(
        self,
        *,
        parsed_contract,
        chunk_count: int,
        pooled_issue_count: int,
        conflicts: list[ConflictRecord],
        decisions: list[ArbitrationDecision],
        final_issues: list[FinalArbitratedIssue],
        elapsed_seconds: float,
    ) -> str:
        """阶段 5：输出仲裁后的最终报告。"""
        result_dir = ensure_multi_agent_demo_result_dir(self.demo_config)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_stem = Path(parsed_contract.file_path).stem
        result_path = result_dir / f"merge_arbitration_{file_stem}_{timestamp}.txt"

        lines: list[str] = [
            "归并层仲裁 Demo 最终报告",
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"源文件路径：{parsed_contract.file_path}",
            f"源文件类型：{parsed_contract.file_type}",
            f"分块数量：{chunk_count}",
            f"原始意见数量：{pooled_issue_count}",
            f"冲突数量：{len(conflicts)}",
            f"仲裁后最终意见数量：{len(final_issues)}",
            f"总计耗时（秒）：{elapsed_seconds:.3f}",
            "",
            "## 仲裁决策",
        ]

        if not decisions:
            lines.append("未触发额外仲裁决策。")
            lines.append("")
        else:
            for decision in decisions:
                lines.append(f"- 决策标题：{decision.title}")
                lines.append(f"  处理动作：{decision.action}")
                lines.append(f"  决策理由：{decision.rationale}")
                lines.append(f"  联动标红分块：{','.join(map(str, decision.highlight_chunks)) or '无'}")
                lines.append("")

        lines.append("## 最终仲裁意见")
        if not final_issues:
            lines.append("未形成最终仲裁意见。")
            lines.append("")
        else:
            for issue in final_issues:
                lines.append(f"### {issue.title}")
                lines.append(f"风险等级：{issue.risk_level}")
                lines.append(f"涉及分块：{','.join(map(str, issue.source_chunks))}")
                lines.append(f"问题说明：{issue.problem}")
                lines.append(f"最终建议：{issue.suggestion}")
                lines.append(f"证据：{issue.evidence}")
                lines.append("")

        with open(result_path, "w", encoding="utf-8") as file:
            file.write("\n".join(lines).strip() + "\n")
        return str(result_path)

    def _extract_json_payload(self, raw_output: str) -> dict[str, Any]:
        """从模型输出中提取 JSON 对象。"""
        text = raw_output.strip()
        if not text:
            raise ValueError("空输出")
        fenced_match = re.search(r"```json\s*(\{[\s\S]*\})\s*```", text)
        if fenced_match:
            text = fenced_match.group(1)
        else:
            object_match = re.search(r"\{[\s\S]*\}", text)
            if object_match:
                text = object_match.group(0)
        payload = json.loads(text)
        if not isinstance(payload, dict):
            raise ValueError("输出不是 JSON 对象")
        return payload

    def _normalize_clause_values(self, raw_values: Any) -> dict[str, str]:
        """清洗关键条款值。"""
        if not isinstance(raw_values, dict):
            return {}
        normalized: dict[str, str] = {}
        for key, value in raw_values.items():
            key_text = str(key).strip()
            value_text = str(value).strip()
            if key_text and value_text:
                normalized[key_text] = value_text
        return normalized

    def _normalize_fields(self, raw_fields: Any) -> list[str]:
        """清洗待比对字段列表。"""
        if not isinstance(raw_fields, list):
            return []
        result: list[str] = []
        for item in raw_fields:
            text = str(item).strip()
            if text and text not in result:
                result.append(text)
        return result

    def _normalize_risk_level(self, value: Any) -> str:
        """标准化风险等级。"""
        text = str(value).strip()
        if text in {"高", "中", "低"}:
            return text
        return "中"

    def _normalize_key(self, value: str) -> str:
        """标准化键名。"""
        return re.sub(r"\s+", "", str(value).strip()).lower()

    def _normalize_value(self, value: str) -> str:
        """标准化字段值。"""
        return re.sub(r"\s+", "", str(value).strip())


class _FakeChatCompletions:
    """用于本地结构验证的假模型。"""

    def create(self, *, model: str, messages: list, tools: list, **kwargs):
        user_prompt = messages[-1]["content"]
        if "10个工作日内付款" in user_prompt:
            content = json.dumps(
                {
                    "chunk_summary": "付款条款要求付款期限与验收条件一起核验。",
                    "issues": [
                        {
                            "topic": "付款期限",
                            "problem": "本块付款期限约定为10个工作日内付款，需要和其他付款条款比对一致性。",
                            "suggestion": "统一为验收完成后10个工作日内付款。",
                            "evidence": "甲方应在验收后10个工作日内付款。",
                            "risk_level": "高",
                            "key_clause_values": {"付款期限": "10个工作日内", "验收条件": "验收后"},
                            "fields_to_compare": ["付款期限", "验收条件"],
                        }
                    ],
                    "key_clause_values": {"付款期限": "10个工作日内", "验收条件": "验收后"},
                    "fields_to_compare": ["付款期限", "验收条件"],
                },
                ensure_ascii=False,
            )
        else:
            content = json.dumps(
                {
                    "chunk_summary": "另一处付款条款给出不同付款期限，需要全局冲突仲裁。",
                    "issues": [
                        {
                            "topic": "付款期限",
                            "problem": "本块付款期限写为15个工作日，与其他分块可能不一致。",
                            "suggestion": "统一为验收完成后15个工作日内付款。",
                            "evidence": "甲方应在验收后15个工作日内付款。",
                            "risk_level": "中",
                            "key_clause_values": {"付款期限": "15个工作日内", "验收条件": "验收后"},
                            "fields_to_compare": ["付款期限", "验收条件"],
                        }
                    ],
                    "key_clause_values": {"付款期限": "15个工作日内", "验收条件": "验收后"},
                    "fields_to_compare": ["付款期限", "验收条件"],
                },
                ensure_ascii=False,
            )
        return SimpleNamespace(
            choices=[
                SimpleNamespace(
                    message=SimpleNamespace(content=content),
                )
            ]
        )


class _FakeLLM:
    """用于本地结构验证的假客户端。"""

    def __init__(self):
        self.chat = SimpleNamespace(completions=_FakeChatCompletions())


async def _main_test_merge_arbitration_demo():
    """本地最小结构验证。"""
    with tempfile.TemporaryDirectory(prefix="merge_arbitration_demo_test_") as temp_dir:
        sample_path = Path(temp_dir) / "sample_contract.docx"
        result_dir = Path(temp_dir) / "result"
        document = Document()
        document.add_paragraph(
            "付款条款一：甲方应在验收后10个工作日内付款。"
            "为避免争议，付款义务应与验收结论绑定，且不得无故拖延。"
            "该条款用于测试跨块字段冲突。"
        )
        document.add_paragraph(
            "付款条款二：甲方应在验收后15个工作日内付款。"
            "如果与前文付款期限不一致，应由主 agent 在全局范围内协调仲裁。"
            "该条款同样用于测试跨块字段冲突。"
        )
        document.save(sample_path)

        demo_config = MultiAgentDemoConfig(
            model=MultiAgentDemoModelConfig(
                model_name="fake-merge-arbitration-model",
                api_key="fake-key",
                api_base="https://example.com/v1",
                temperature=0.0,
                top_p=1.0,
                max_tokens=1024,
            ),
            chunk_size=120,
            max_concurrent_reviews=2,
            result_dir=str(result_dir),
        )
        service = MergeArbitrationReviewDemo(
            config=demo_config,
            llm_client=_FakeLLM(),
        )

        result = await service.run(str(sample_path))
        print("merge_arbitration_demo self test result:")
        print(
            {
                "chunk_count": result["chunk_count"],
                "conflict_count": result["conflict_count"],
                "final_issue_count": result["final_issue_count"],
                "result_file_exists": Path(result["result_file_path"]).exists(),
            }
        )
        assert result["chunk_count"] >= 2
        assert result["conflict_count"] >= 1
        assert result["final_issue_count"] >= 1
        assert Path(result["result_file_path"]).exists()
        print("merge_arbitration_demo self test passed")


if __name__ == "__main__":
    asyncio.run(_main_test_merge_arbitration_demo())
