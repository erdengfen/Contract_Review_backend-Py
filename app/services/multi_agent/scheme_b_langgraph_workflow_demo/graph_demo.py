"""方案 B：LangGraph Workflow Demo。

多 agent 链路：
1. parse_document
2. split_and_index
3. router
4. specialist_payment / specialist_liability / specialist_dispute / specialist_definition
5. merge_evidence
6. rule_gate_1
7. arbitration
8. rule_gate_2
9. record_disputed_findings
10. reporter

说明：
- 本文件直接使用 LangGraph 实现最小图骨架。
- 当前阶段只实现 Phase B1 最小闭环。
- 数据库与 RAG 仅保留字段，不做真实接入。
"""

from __future__ import annotations

import argparse
import asyncio
import json
import tempfile
import time
from collections import Counter
from datetime import datetime
from pathlib import Path
from types import SimpleNamespace
from typing import Any

from docx import Document
from langgraph.graph import END, START, StateGraph

try:
    from ..config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        ensure_multi_agent_demo_result_dir,
        get_multi_agent_demo_config,
        init_multi_agent_demo_llm,
    )
    from .graph_state import GraphState
    from .node_handlers import SchemeBLangGraphNodeHandlers
except ImportError:
    from app.services.multi_agent.config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        ensure_multi_agent_demo_result_dir,
        get_multi_agent_demo_config,
        init_multi_agent_demo_llm,
    )
    from app.services.multi_agent.scheme_b_langgraph_workflow_demo.graph_state import (
        GraphState,
    )
    from app.services.multi_agent.scheme_b_langgraph_workflow_demo.node_handlers import (
        SchemeBLangGraphNodeHandlers,
    )


class SchemeBLangGraphWorkflowDemo:
    """方案 B demo 主入口。"""

    def __init__(
        self,
        config: MultiAgentDemoConfig | None = None,
        llm_client: Any | None = None,
    ):
        self.demo_config = config or get_multi_agent_demo_config()
        self.llm = llm_client or init_multi_agent_demo_llm(self.demo_config.model)
        self.handlers = SchemeBLangGraphNodeHandlers(
            config=self.demo_config,
            llm_client=self.llm,
        )
        self.graph = self._build_graph()

    def _build_graph(self):
        """构建 LangGraph 主图。"""
        graph = StateGraph(GraphState)
        graph.add_node("parse_document", self.handlers.parse_document_node)
        graph.add_node("split_and_index", self.handlers.split_and_index_node)
        graph.add_node("router", self.handlers.router_node)
        graph.add_node("specialist_payment", self.handlers.specialist_payment_node)
        graph.add_node("specialist_liability", self.handlers.specialist_liability_node)
        graph.add_node("specialist_dispute", self.handlers.specialist_dispute_node)
        graph.add_node("specialist_definition", self.handlers.specialist_definition_node)
        graph.add_node("merge_evidence", self.handlers.merge_evidence_node)
        graph.add_node("rule_gate_1", self.handlers.rule_gate_1_node)
        graph.add_node("arbitration", self.handlers.arbitration_node)
        graph.add_node("rule_gate_2", self.handlers.rule_gate_2_node)
        graph.add_node(
            "record_disputed_findings",
            self.handlers.record_disputed_findings_node,
        )
        graph.add_node("reporter", self._reporter_node)

        graph.add_edge(START, "parse_document")
        graph.add_edge("parse_document", "split_and_index")
        graph.add_edge("split_and_index", "router")
        graph.add_edge("router", "specialist_payment")
        graph.add_edge("router", "specialist_liability")
        graph.add_edge("router", "specialist_dispute")
        graph.add_edge("router", "specialist_definition")
        graph.add_edge("specialist_payment", "merge_evidence")
        graph.add_edge("specialist_liability", "merge_evidence")
        graph.add_edge("specialist_dispute", "merge_evidence")
        graph.add_edge("specialist_definition", "merge_evidence")
        graph.add_edge("merge_evidence", "rule_gate_1")
        graph.add_edge("rule_gate_1", "arbitration")
        graph.add_edge("arbitration", "rule_gate_2")
        graph.add_edge("rule_gate_2", "record_disputed_findings")
        graph.add_edge("record_disputed_findings", "reporter")
        graph.add_edge("reporter", END)
        return graph.compile()

    async def run(self, file_path: str) -> dict[str, Any]:
        """运行方案 B demo。"""
        start_time = time.perf_counter()
        run_id = f"scheme-b-{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        final_state = await self.graph.ainvoke(
            {
                "run_id": run_id,
                "file_path": str(Path(file_path).expanduser().resolve()),
                "execution_trace": [],
                "specialist_outputs": [],
            }
        )
        final_state["elapsed_seconds"] = time.perf_counter() - start_time
        return final_state

    async def _reporter_node(self, state: GraphState) -> dict[str, Any]:
        """结果落盘。"""
        started = time.perf_counter()
        result_dir = ensure_multi_agent_demo_result_dir(self.demo_config)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_stem = Path(state["file_path"]).stem
        text_path = result_dir / f"scheme_b_langgraph_{file_stem}_{timestamp}.txt"
        json_path = result_dir / f"scheme_b_langgraph_{file_stem}_{timestamp}.json"

        specialist_counter = Counter(
            item["specialist_name"] for item in state.get("specialist_outputs", [])
        )

        lines: list[str] = [
            "方案 B：LangGraph Workflow Demo 运行结果",
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"源文件路径：{state['file_path']}",
            f"源文件类型：{state.get('file_type', 'unknown')}",
            f"运行 ID：{state['run_id']}",
            f"证据快照版本：{state.get('snapshot_id', '')}",
            f"分块数量：{len(state.get('chunks', []))}",
            f"路由任务数量：{len(state.get('router_tasks', []))}",
            f"specialist 输出数量：{len(state.get('specialist_outputs', []))}",
            f"merged findings 数量：{len(state.get('merged_findings', []))}",
            f"accepted 数量：{len(state.get('accepted_findings', []))}",
            f"disputed 数量：{len(state.get('disputed_findings', []))}",
            f"suppressed 数量：{len(state.get('suppressed_findings', []))}",
            "",
            "## Specialist 分布",
        ]
        for name, count in sorted(specialist_counter.items()):
            lines.append(f"- {name}：{count}")

        lines.extend(["", "## 路由统计"])
        routing_summary = state.get("routing_summary", {})
        if routing_summary:
            lines.append(f"- task_count：{routing_summary.get('task_count', 0)}")
            lines.append(
                f"- topic_distribution：{json.dumps(routing_summary.get('topic_distribution', {}), ensure_ascii=False)}"
            )
            lines.append(
                f"- specialist_distribution：{json.dumps(routing_summary.get('specialist_distribution', {}), ensure_ascii=False)}"
            )

        lines.extend(["", "## 规则统计"])
        rule_summary = state.get("rule_summary", {})
        if rule_summary:
            lines.append(f"- disputed_count：{rule_summary.get('disputed_count', 0)}")
            lines.append(
                f"- gate_distribution：{json.dumps(rule_summary.get('gate_distribution', {}), ensure_ascii=False)}"
            )
            lines.append(
                f"- rule_distribution：{json.dumps(rule_summary.get('rule_distribution', {}), ensure_ascii=False)}"
            )

        lines.extend(["", "## 仲裁摘要"])
        decisions = state.get("arbitration_decisions", [])
        if not decisions:
            lines.append("无")
        else:
            for item in decisions:
                lines.append(f"- {item['title']}：{item['action']}，{item['rationale']}")

        lines.extend(["", "## Disputed Findings 摘要"])
        disputed_findings = state.get("disputed_findings", [])
        if not disputed_findings:
            lines.append("无")
        else:
            for item in disputed_findings:
                rule_ids = [hit["rule_id"] for hit in item["rule_hits"]]
                lines.append(
                    f"- [{item['gate_name']}] {item['title']} | 风险={item['risk_level']} | 规则={','.join(rule_ids) or '无'} | 原因={item['dispute_reason']}"
                )

        lines.extend(["", "## 执行轨迹"])
        for item in state.get("execution_trace", []):
            lines.append(
                f"- {item['node_name']}：{item['status']}，耗时 {item['elapsed_ms']}ms，{item['summary']}"
            )

        with open(text_path, "w", encoding="utf-8") as file:
            file.write("\n".join(lines).strip() + "\n")

        payload = dict(state)
        payload["elapsed_seconds"] = payload.get("elapsed_seconds", 0.0) + (
            time.perf_counter() - started
        )
        with open(json_path, "w", encoding="utf-8") as file:
            json.dump(payload, file, ensure_ascii=False, indent=2)

        return {
            "result_file_path": str(text_path),
            "result_json_path": str(json_path),
            "execution_trace": [
                {
                    "trace_id": f"trace-reporter-{timestamp}",
                    "node_name": "reporter",
                    "status": "completed",
                    "started_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "ended_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "elapsed_ms": int((time.perf_counter() - started) * 1000),
                    "summary": "文本与 JSON 结果已落盘",
                    "input_count": len(state.get("accepted_findings", [])) + len(state.get("disputed_findings", [])),
                    "output_count": 2,
                }
            ],
        }


class _FakeChatCompletions:
    """用于结构验证的假模型。"""

    def create(self, *, model: str, messages: list, tools: list, **kwargs):
        user_prompt = messages[-1]["content"]
        if "付款 specialist" in user_prompt and "验收合格后30个工作日内支付" in user_prompt:
            content = """
【修改点1】付款期限偏长
【原文】甲方应在验收合格后30个工作日内支付全部合同价款。
【风险分析】付款期限偏长且触发条件较弱，容易引发结算争议。
【风险等级】高
【修改后的内容】甲方应在验收合格并收到合规发票后10个工作日内支付全部合同价款。
【修改理由】缩短付款周期并补足触发条件。
【风险类型】付款机制歧义
""".strip()
        elif "违约 specialist" in user_prompt and "违约金比例留空" in user_prompt:
            content = """
【修改点1】违约责任缺失
【原文】违约责任：任何一方违约均应承担损失，违约金比例留空。
【风险分析】违约金比例为空，责任标准难以执行。
【风险等级】高
【修改后的内容】违约金比例建议明确为合同价款的5%。
【修改理由】补足关键责任字段。
【风险类型】违约责任缺失
""".strip()
        elif "争议 specialist" in user_prompt and "项目所在地法院" in user_prompt:
            content = """
【修改点1】争议解决条款不完整
【原文】争议提交项目所在地法院。
【风险分析】未明确前置协商程序和具体管辖表达。
【风险等级】中
【修改后的内容】争议应先协商解决，协商不成提交项目所在地有管辖权的人民法院。
【修改理由】补足程序链路。
【风险类型】争议解决
""".strip()
        else:
            content = """
【修改点1】定义条款模糊
【原文】本合同所称验收合格是指甲方书面确认。
【风险分析】定义条款较短，标准不够完整。
【风险等级】中
【修改后的内容】验收合格是指甲方依据合同约定标准完成书面确认。
【修改理由】补足定义口径。
【风险类型】定义条款
""".strip()
        return SimpleNamespace(
            choices=[SimpleNamespace(message=SimpleNamespace(content=content))]
        )


class _FakeLLM:
    """用于结构验证的假客户端。"""

    def __init__(self):
        self.chat = SimpleNamespace(completions=_FakeChatCompletions())


async def _main_test_scheme_b_demo():
    """方案 B 最小结构验证。"""
    with tempfile.TemporaryDirectory(prefix="scheme_b_langgraph_test_") as temp_dir:
        sample_path = Path(temp_dir) / "sample_contract.docx"
        result_dir = Path(temp_dir) / "result"

        document = Document()
        document.add_paragraph("定义条款：本合同所称验收合格是指甲方书面确认。")
        document.add_paragraph("付款条款：甲方应在验收合格后30个工作日内支付全部合同价款。")
        document.add_paragraph("违约责任：任何一方违约均应承担损失，违约金比例留空。争议提交项目所在地法院。")
        document.save(sample_path)

        demo_config = MultiAgentDemoConfig(
            model=MultiAgentDemoModelConfig(
                model_name="fake-scheme-b-model",
                api_key="fake-key",
                api_base="https://example.com/v1",
                temperature=0.0,
                top_p=1.0,
                max_tokens=1024,
            ),
            chunk_size=40,
            max_concurrent_reviews=2,
            result_dir=str(result_dir),
        )
        demo = SchemeBLangGraphWorkflowDemo(
            config=demo_config,
            llm_client=_FakeLLM(),
        )
        result = await demo.run(str(sample_path))
        print("scheme_b_langgraph_demo self test result:")
        print(
            {
                "chunk_count": len(result.get("chunks", [])),
                "router_task_count": len(result.get("router_tasks", [])),
                "accepted_count": len(result.get("accepted_findings", [])),
                "disputed_count": len(result.get("disputed_findings", [])),
                "result_file_exists": Path(result["result_file_path"]).exists(),
                "result_json_exists": Path(result["result_json_path"]).exists(),
                "trace_count": len(result.get("execution_trace", [])),
            }
        )
        assert len(result.get("chunks", [])) >= 3
        assert len(result.get("router_tasks", [])) >= 3
        assert Path(result["result_file_path"]).exists()
        assert Path(result["result_json_path"]).exists()
        assert len(result.get("execution_trace", [])) >= 5
        print("scheme_b_langgraph_demo self test passed")


async def _main_real(file_path: str):
    """真实运行入口。"""
    demo = SchemeBLangGraphWorkflowDemo()
    result = await demo.run(file_path)
    print("scheme_b_langgraph_demo real run result:")
    print(
        {
            "chunk_count": len(result.get("chunks", [])),
            "router_task_count": len(result.get("router_tasks", [])),
            "accepted_count": len(result.get("accepted_findings", [])),
            "disputed_count": len(result.get("disputed_findings", [])),
            "suppressed_count": len(result.get("suppressed_findings", [])),
            "elapsed_seconds": round(result.get("elapsed_seconds", 0.0), 3),
            "result_file_path": result.get("result_file_path"),
            "result_json_path": result.get("result_json_path"),
        }
    )


def main():
    """CLI 入口。"""
    parser = argparse.ArgumentParser(description="方案 B LangGraph Workflow Demo")
    parser.add_argument("--file", type=str, help="真实合同文件路径")
    args = parser.parse_args()
    if args.file:
        asyncio.run(_main_real(args.file))
        return
    asyncio.run(_main_test_scheme_b_demo())


if __name__ == "__main__":
    main()
