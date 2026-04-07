"""方案 A：归并仲裁 + 规则引擎后置校验 Demo。

多 agent 链路：
1. 合同解析与分块
2. 构建证据快照
3. Planner 按分块派发 Specialist 任务
4. Specialist 并发审阅
5. 结构化 Findings 汇总
6. Merge + Arbitration 收束冲突
7. Rule Checker 校验
8. 输出 Accepted / Needs Recheck / Suppressed
9. 生成审计报告并落盘

说明：
- 本文件为 demo 骨架，优先复用现有 multi_agent 工具与归并逻辑。
- 当前不接入真实数据库与完整规则 DSL，仅保留关键字段与运行闭环。
"""

from __future__ import annotations

import argparse
import asyncio
import json
import re
import tempfile
import time
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass, field
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path
from types import SimpleNamespace
from typing import Any

from docx import Document

try:
    from ..config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        ensure_multi_agent_demo_result_dir,
        get_multi_agent_demo_config,
    )
    from ..merge_arbitration_demo import (
        ArbitrationDecision,
        ChunkAnalysisOpinion,
        ChunkIssueOpinion,
        ConflictRecord,
        FinalArbitratedIssue,
        MergeArbitrationReviewDemo,
        RISK_SCORE_MAP,
    )
except ImportError:
    from app.services.multi_agent.config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        ensure_multi_agent_demo_result_dir,
        get_multi_agent_demo_config,
    )
    from app.services.multi_agent.merge_arbitration_demo import (
        ArbitrationDecision,
        ChunkAnalysisOpinion,
        ChunkIssueOpinion,
        ConflictRecord,
        FinalArbitratedIssue,
        MergeArbitrationReviewDemo,
        RISK_SCORE_MAP,
    )


REFERENCE_PATTERN = re.compile(
    r"(第[一二三四五六七八九十0-9]+条|附件[一二三四五六七八九十0-9A-Za-z]+|附表[一二三四五六七八九十0-9A-Za-z]+)"
)
PLACEHOLDER_TOKENS = (
    "本块原文中的问题",
    "修改建议",
    "未找到修改建议",
)
GENERIC_EVIDENCE_TOKENS = (
    "多个分块提取的关键字段值不一致",
    "部分修改建议会直接影响其他分块中的对应字段",
    "结构化解析失败",
    "未形成有效证据",
)
GENERIC_SUGGESTION_TOKENS = (
    "建议进一步明确",
    "建议补充约定",
    "建议结合实际情况",
    "建议双方协商",
    "建议评估",
    "建议审慎",
    "建议关注",
    "建议核实",
    "修改建议",
)
SUMMARY_TITLE_TOKENS = (
    "总结",
    "行动建议",
    "核心风险提示",
)
TITLE_CANONICALIZE_TOKENS = (
    "风险",
    "条款",
    "约定",
    "不明",
    "模糊",
    "缺失",
    "不足",
    "不清",
    "不完整",
    "过宽",
    "过高",
    "过低",
    "过严",
    "过苛",
    "绝对化",
    "绝对",
    "瑕疵",
    "不合理",
    "不当",
    "失衡",
    "缺陷",
    "薄弱",
)


@dataclass(slots=True)
class SpecialistProfile:
    """Specialist 角色配置。"""

    specialist_name: str
    focus_area: str
    instruction: str
    keywords: list[str] = field(default_factory=list)
    compare_fields: list[str] = field(default_factory=list)


SPECIALIST_PROFILES = [
    SpecialistProfile(
        specialist_name="付款与结算 specialist",
        focus_area="付款与结算",
        instruction="重点审阅付款触发条件、付款期限、结算口径、预付款与验收联动。",
        keywords=["付款", "支付", "结算", "价款", "预付款", "进度款", "验收"],
        compare_fields=["付款期限", "验收条件"],
    ),
    SpecialistProfile(
        specialist_name="责任与违约 specialist",
        focus_area="责任与违约",
        instruction="重点审阅违约责任、赔偿范围、违约金比例、责任限制与索赔条款。",
        keywords=["责任", "违约", "赔偿", "违约金", "索赔", "免责"],
        compare_fields=["违约金比例", "赔偿上限"],
    ),
    SpecialistProfile(
        specialist_name="终止与争议 specialist",
        focus_area="终止与争议",
        instruction="重点审阅解除、终止、争议解决、通知、适用法律与管辖条款。",
        keywords=["解除", "终止", "争议", "仲裁", "法院", "通知", "法律"],
        compare_fields=["争议解决地", "通知地址"],
    ),
    SpecialistProfile(
        specialist_name="定义与异常 specialist",
        focus_area="定义与异常",
        instruction="重点审阅定义项、附件回指、空白字段、占位内容与格式异常。",
        keywords=["定义", "系指", "是指", "附件", "附表", "空白", "留空"],
        compare_fields=["附件编号", "定义项"],
    ),
]


@dataclass(slots=True)
class EvidenceChunk:
    """证据快照中的分块节点。"""

    chunk_id: str
    chunk_index: int
    text: str
    prev_chunk_id: str = ""
    next_chunk_id: str = ""
    keyword_tags: list[str] = field(default_factory=list)
    reference_tokens: list[str] = field(default_factory=list)


@dataclass(slots=True)
class DefinitionRecord:
    """轻量定义项记录。"""

    definition_id: str
    term: str
    description: str
    source_chunk_id: str


@dataclass(slots=True)
class EvidenceSnapshot:
    """共享证据快照。"""

    snapshot_id: str
    file_path: str
    file_type: str
    chunks: list[EvidenceChunk] = field(default_factory=list)
    definitions: list[DefinitionRecord] = field(default_factory=list)
    available_handles: list[str] = field(default_factory=list)


@dataclass(slots=True)
class ContextPackage:
    """Specialist 任务上下文包。"""

    chunk_id: str
    current_text: str
    adjacent_texts: list[str] = field(default_factory=list)
    related_definition_texts: list[str] = field(default_factory=list)
    reference_tokens: list[str] = field(default_factory=list)
    available_handles: list[str] = field(default_factory=list)


@dataclass(slots=True)
class SpecialistTask:
    """Planner 派发的任务。"""

    task_id: str
    chunk_id: str
    chunk_index: int
    specialist_name: str
    focus_area: str
    instruction: str
    compare_fields: list[str] = field(default_factory=list)
    context_package: ContextPackage | None = None


@dataclass(slots=True)
class RuleHit:
    """规则命中记录。"""

    rule_id: str
    priority: str
    result: str
    message: str


@dataclass(slots=True)
class CheckedFinding:
    """规则校验后的最终发现。"""

    finding_id: str
    title: str
    risk_level: str
    risk_score: int
    issue: str
    suggestion: str
    evidence: str
    evidence_chunk_ids: list[str] = field(default_factory=list)
    dependency_ids: list[str] = field(default_factory=list)
    source_chunk_ids: list[str] = field(default_factory=list)
    source_issue_ids: list[str] = field(default_factory=list)
    checker_status: str = "accepted"
    checker_notes: list[str] = field(default_factory=list)
    rule_hits: list[RuleHit] = field(default_factory=list)


@dataclass(slots=True)
class StageTiming:
    """阶段耗时记录。"""

    stage_name: str
    elapsed_seconds: float


class MergeRuleEngineReviewDemo(MergeArbitrationReviewDemo):
    """方案 A 最小可运行骨架。"""

    max_specialists_per_chunk = 2
    default_specialists_per_chunk = 1
    secondary_specialist_min_score = 4
    secondary_specialist_score_gap = 1

    async def run(
        self,
        file_path: str,
        *,
        stance: str | None = None,
        intensity: str | None = None,
        contract_type: str | None = None,
    ) -> dict[str, Any]:
        """执行方案 A demo。"""
        start_time = time.perf_counter()
        stage_timings: list[StageTiming] = []

        stage_started = time.perf_counter()
        parsed_contract = await self.parse_contract_file(file_path)
        chunks = self.split_contract_text(parsed_contract.text)
        stage_timings.append(
            StageTiming(
                stage_name="parse_and_split",
                elapsed_seconds=time.perf_counter() - stage_started,
            )
        )

        stage_started = time.perf_counter()
        snapshot = self._build_evidence_snapshot(
            file_path=parsed_contract.file_path,
            file_type=parsed_contract.file_type,
            chunks=chunks,
        )
        stage_timings.append(
            StageTiming(
                stage_name="build_snapshot",
                elapsed_seconds=time.perf_counter() - stage_started,
            )
        )

        stage_started = time.perf_counter()
        tasks = self._plan_specialist_tasks(snapshot)
        stage_timings.append(
            StageTiming(
                stage_name="planner_dispatch",
                elapsed_seconds=time.perf_counter() - stage_started,
            )
        )

        stage_started = time.perf_counter()
        chunk_opinions = await self._run_specialist_reviews(
            snapshot=snapshot,
            tasks=tasks,
            stance=stance or self.demo_config.default_stance,
            intensity=intensity or self.demo_config.default_intensity,
            contract_type=contract_type or self.demo_config.default_contract_type,
        )
        stage_timings.append(
            StageTiming(
                stage_name="specialist_reviews",
                elapsed_seconds=time.perf_counter() - stage_started,
            )
        )

        stage_started = time.perf_counter()
        pooled_issues = self._collect_global_opinion_pool(chunk_opinions)
        conflicts = self._detect_conflicts(chunk_opinions, pooled_issues)
        decisions, final_issues = self._arbitrate_conflicts(
            chunk_opinions=chunk_opinions,
            pooled_issues=pooled_issues,
            conflicts=conflicts,
        )
        stage_timings.append(
            StageTiming(
                stage_name="merge_and_arbitration",
                elapsed_seconds=time.perf_counter() - stage_started,
            )
        )

        stage_started = time.perf_counter()
        final_issues, precheck_merge_count = self._precheck_consolidate_final_issues(
            final_issues
        )
        stage_timings.append(
            StageTiming(
                stage_name="precheck_consolidation",
                elapsed_seconds=time.perf_counter() - stage_started,
            )
        )

        stage_started = time.perf_counter()
        accepted_findings, needs_recheck_findings, suppressed_findings = (
            self._run_rule_checker(final_issues)
        )
        stage_timings.append(
            StageTiming(
                stage_name="rule_checker",
                elapsed_seconds=time.perf_counter() - stage_started,
            )
        )

        elapsed_seconds_without_save = time.perf_counter() - start_time
        stage_started = time.perf_counter()
        result_paths = self._save_run_result(
            snapshot=snapshot,
            tasks=tasks,
            chunk_opinions=chunk_opinions,
            conflicts=conflicts,
            decisions=decisions,
            precheck_merge_count=precheck_merge_count,
            accepted_findings=accepted_findings,
            needs_recheck_findings=needs_recheck_findings,
            suppressed_findings=suppressed_findings,
            elapsed_seconds=elapsed_seconds_without_save,
            stage_timings=stage_timings,
        )
        stage_timings.append(
            StageTiming(
                stage_name="save_result",
                elapsed_seconds=time.perf_counter() - stage_started,
            )
        )
        elapsed_seconds = time.perf_counter() - start_time
        return {
            "demo_type": "scheme_a_merge_rule_engine",
            "snapshot_id": snapshot.snapshot_id,
            "chunk_count": len(snapshot.chunks),
            "task_count": len(tasks),
            "pooled_issue_count": len(pooled_issues),
            "conflict_count": len(conflicts),
            "precheck_merge_count": precheck_merge_count,
            "accepted_count": len(accepted_findings),
            "needs_recheck_count": len(needs_recheck_findings),
            "suppressed_count": len(suppressed_findings),
            "elapsed_seconds": elapsed_seconds,
            "stage_timings": [asdict(item) for item in stage_timings],
            "result_file_path": result_paths["text"],
            "result_json_path": result_paths["json"],
        }

    def _build_evidence_snapshot(
        self,
        *,
        file_path: str,
        file_type: str,
        chunks: list[str],
    ) -> EvidenceSnapshot:
        """构建共享证据快照。"""
        snapshot_chunks: list[EvidenceChunk] = []
        definitions: list[DefinitionRecord] = []

        for index, chunk_text in enumerate(chunks, start=1):
            chunk_id = f"chunk-{index}"
            snapshot_chunks.append(
                EvidenceChunk(
                    chunk_id=chunk_id,
                    chunk_index=index,
                    text=chunk_text,
                    prev_chunk_id=f"chunk-{index - 1}" if index > 1 else "",
                    next_chunk_id=f"chunk-{index + 1}" if index < len(chunks) else "",
                    keyword_tags=self._infer_chunk_tags(chunk_text),
                    reference_tokens=self._extract_reference_tokens(chunk_text),
                )
            )
            for term, description in self._extract_definition_pairs(chunk_text):
                definitions.append(
                    DefinitionRecord(
                        definition_id=f"definition-{len(definitions) + 1}",
                        term=term,
                        description=description,
                        source_chunk_id=chunk_id,
                    )
                )

        available_handles = [item.chunk_id for item in snapshot_chunks]
        available_handles.extend(item.definition_id for item in definitions)
        available_handles.extend(
            token for item in snapshot_chunks for token in item.reference_tokens
        )
        return EvidenceSnapshot(
            snapshot_id=f"snapshot-{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            file_path=file_path,
            file_type=file_type,
            chunks=snapshot_chunks,
            definitions=definitions,
            available_handles=self._unique_texts(available_handles),
        )

    def _plan_specialist_tasks(
        self,
        snapshot: EvidenceSnapshot,
    ) -> list[SpecialistTask]:
        """Planner 根据证据快照派发 Specialist 任务。"""
        tasks: list[SpecialistTask] = []
        for chunk in snapshot.chunks:
            context_package = self._build_context_package(snapshot, chunk)
            scored_profiles = self._score_specialist_profiles(chunk)
            selected_profiles = self._select_specialist_profiles(
                chunk=chunk,
                scored_profiles=scored_profiles,
            )

            for profile in selected_profiles:
                task_id = f"{chunk.chunk_id}-{self._normalize_key(profile.focus_area)}"
                tasks.append(
                    SpecialistTask(
                        task_id=task_id,
                        chunk_id=chunk.chunk_id,
                        chunk_index=chunk.chunk_index,
                        specialist_name=profile.specialist_name,
                        focus_area=profile.focus_area,
                        instruction=profile.instruction,
                        compare_fields=profile.compare_fields,
                        context_package=context_package,
                    )
                )
        return tasks

    def _score_specialist_profiles(
        self,
        chunk: EvidenceChunk,
    ) -> list[tuple[int, SpecialistProfile]]:
        """按 chunk 信号为 specialist 打分。"""
        scored_profiles: list[tuple[int, SpecialistProfile]] = []
        placeholder_signal = self._contains_placeholder_signal(chunk.text)
        compare_values, compare_fields = self._extract_chunk_compare_metadata(chunk.text)
        compare_blob = " ".join(compare_fields + list(compare_values.keys()))

        for profile in SPECIALIST_PROFILES:
            matched_keywords = [
                keyword for keyword in profile.keywords if keyword in chunk.text
            ]
            score = len(matched_keywords) * 2

            for field_name in profile.compare_fields:
                if field_name and field_name in compare_blob:
                    score += 2

            if profile.focus_area == "定义与异常":
                if chunk.reference_tokens:
                    score += 2
                if placeholder_signal:
                    score += 3
                if "附件" in chunk.text or "附表" in chunk.text:
                    score += 1

            if profile.focus_area == "终止与争议" and any(
                token in chunk.text for token in ("争议", "仲裁", "法院", "通知")
            ):
                score += 1

            if profile.focus_area == "责任与违约" and any(
                token in chunk.text for token in ("违约", "赔偿", "责任限制", "免责")
            ):
                score += 1

            if profile.focus_area == "付款与结算" and any(
                token in chunk.text for token in ("价款", "支付条件", "进度款", "结算")
            ):
                score += 1

            if score > 0:
                scored_profiles.append((score, profile))

        scored_profiles.sort(key=lambda item: (-item[0], item[1].specialist_name))
        return scored_profiles

    def _select_specialist_profiles(
        self,
        *,
        chunk: EvidenceChunk,
        scored_profiles: list[tuple[int, SpecialistProfile]],
    ) -> list[SpecialistProfile]:
        """收紧 Planner 派单策略，避免每块固定扩成双 specialist。"""
        if not scored_profiles:
            if self._should_skip_chunk_dispatch(chunk):
                return []
            return [SPECIALIST_PROFILES[-1]]

        selected_profiles = [
            item[1] for item in scored_profiles[: self.default_specialists_per_chunk]
        ]
        if len(scored_profiles) <= 1:
            return selected_profiles

        if not self._should_add_secondary_specialist(chunk, scored_profiles):
            return selected_profiles

        selected_profiles.append(scored_profiles[1][1])
        return selected_profiles[: self.max_specialists_per_chunk]

    def _should_skip_chunk_dispatch(self, chunk: EvidenceChunk) -> bool:
        """低信号 chunk 不派发 specialist，由后续全局联动补位。"""
        if chunk.reference_tokens:
            return False
        if self._contains_placeholder_signal(chunk.text):
            return False
        compare_values, compare_fields = self._extract_chunk_compare_metadata(chunk.text)
        if compare_values or compare_fields:
            return False
        if len(chunk.keyword_tags) >= 2:
            return False
        return len(self._normalize_free_text(chunk.text)) < 220

    def _should_add_secondary_specialist(
        self,
        chunk: EvidenceChunk,
        scored_profiles: list[tuple[int, SpecialistProfile]],
    ) -> bool:
        """仅在多焦点且强信号分块中追加第二个 specialist。"""
        primary_score = scored_profiles[0][0]
        secondary_score = scored_profiles[1][0]
        if secondary_score < self.secondary_specialist_min_score:
            return False
        if primary_score - secondary_score > self.secondary_specialist_score_gap:
            return False

        compare_values, compare_fields = self._extract_chunk_compare_metadata(chunk.text)
        strong_signal_count = 0
        if len(chunk.reference_tokens) >= 1:
            strong_signal_count += 1
        if len(compare_values) >= 1 or len(compare_fields) >= 1:
            strong_signal_count += 1
        if len(chunk.keyword_tags) >= 4:
            strong_signal_count += 1
        if len(self._normalize_free_text(chunk.text)) >= 360:
            strong_signal_count += 1
        if self._contains_placeholder_signal(chunk.text):
            strong_signal_count += 1
        return strong_signal_count >= 2

    def _build_context_package(
        self,
        snapshot: EvidenceSnapshot,
        chunk: EvidenceChunk,
    ) -> ContextPackage:
        """组装 Specialist 的最小上下文包。"""
        chunk_map = {item.chunk_id: item for item in snapshot.chunks}
        adjacent_texts: list[str] = []
        for chunk_id in (chunk.prev_chunk_id, chunk.next_chunk_id):
            if chunk_id and chunk_id in chunk_map:
                adjacent_texts.append(self._truncate_text(chunk_map[chunk_id].text))

        related_definition_texts: list[str] = []
        for definition in snapshot.definitions:
            if definition.term and definition.term in chunk.text:
                related_definition_texts.append(
                    f"{definition.term}：{self._truncate_text(definition.description)}"
                )

        available_handles = [chunk.chunk_id]
        available_handles.extend(chunk.reference_tokens)
        available_handles.extend(
            definition.definition_id
            for definition in snapshot.definitions
            if definition.term and definition.term in chunk.text
        )

        return ContextPackage(
            chunk_id=chunk.chunk_id,
            current_text=chunk.text,
            adjacent_texts=adjacent_texts,
            related_definition_texts=related_definition_texts,
            reference_tokens=chunk.reference_tokens,
            available_handles=self._unique_texts(available_handles),
        )

    async def _run_specialist_reviews(
        self,
        *,
        snapshot: EvidenceSnapshot,
        tasks: list[SpecialistTask],
        stance: str,
        intensity: str,
        contract_type: str,
    ) -> list[ChunkAnalysisOpinion]:
        """并发执行 Specialist 审阅，并按 chunk 汇总。"""
        semaphore = asyncio.Semaphore(self.demo_config.max_concurrent_reviews)
        ordered_task_results: list[tuple[SpecialistTask, list[ChunkIssueOpinion]] | None] = [
            None
        ] * len(tasks)

        async def _run_single(index: int, task: SpecialistTask):
            async with semaphore:
                issues = await self._run_single_specialist_task(
                    task=task,
                    stance=stance,
                    intensity=intensity,
                    contract_type=contract_type,
                )
                ordered_task_results[index] = (task, issues)

        await asyncio.gather(
            *[_run_single(index, task) for index, task in enumerate(tasks)]
        )

        chunk_map = {item.chunk_id: item for item in snapshot.chunks}
        grouped_issues: dict[str, list[ChunkIssueOpinion]] = defaultdict(list)
        grouped_specialists: dict[str, list[str]] = defaultdict(list)
        for item in ordered_task_results:
            if item is None:
                continue
            task, issues = item
            grouped_issues[task.chunk_id].extend(issues)
            grouped_specialists[task.chunk_id].append(task.specialist_name)

        chunk_opinions: list[ChunkAnalysisOpinion] = []
        for chunk in snapshot.chunks:
            key_values, compare_fields = self._extract_chunk_compare_metadata(
                chunk.text
            )
            chunk_opinions.append(
                ChunkAnalysisOpinion(
                    chunk_index=chunk.chunk_index,
                    chunk_text=chunk.text,
                    chunk_summary=(
                        "参与 specialist："
                        + (
                            "、".join(
                                self._unique_texts(grouped_specialists[chunk.chunk_id])
                            )
                            or "未派发，等待全局归并与规则复核"
                        )
                    ),
                    issues=grouped_issues.get(chunk.chunk_id, []),
                    key_clause_values=key_values,
                    fields_to_compare=compare_fields,
                    raw_output="",
                )
            )
        return chunk_opinions

    async def _run_single_specialist_task(
        self,
        *,
        task: SpecialistTask,
        stance: str,
        intensity: str,
        contract_type: str,
    ) -> list[ChunkIssueOpinion]:
        """执行单个 Specialist 任务。"""
        context_package = task.context_package or ContextPackage(
            chunk_id=task.chunk_id,
            current_text="",
        )
        context = (
            f"你当前扮演 {task.specialist_name}。\n"
            f"审阅焦点：{task.focus_area}。\n"
            f"执行指令：{task.instruction}\n"
            f"当前分块：{task.chunk_id}\n"
            f"相邻窗口：{'；'.join(context_package.adjacent_texts) or '无'}\n"
            f"相关定义：{'；'.join(context_package.related_definition_texts) or '无'}\n"
            f"引用标记：{'；'.join(context_package.reference_tokens) or '无'}\n"
            f"可用句柄：{'；'.join(context_package.available_handles) or '无'}"
        )
        review_result = await self.review_chunk(
            chunk_text=context_package.current_text,
            stance=stance,
            intensity=intensity,
            contract_type=contract_type,
            context=context,
        )
        key_values, extracted_fields = self._extract_chunk_compare_metadata(
            context_package.current_text
        )
        issues: list[ChunkIssueOpinion] = []
        for index, item in enumerate(review_result, start=1):
            fields_to_compare = self._unique_texts(
                extracted_fields + task.compare_fields
            )
            issues.append(
                ChunkIssueOpinion(
                    issue_id=f"{task.task_id}-issue-{index}",
                    topic=str(item.get("risk_type", "")).strip() or task.focus_area,
                    problem=str(item.get("risk_analysis", "")).strip(),
                    suggestion=str(item.get("suggested_content", "")).strip(),
                    evidence=str(item.get("original_content", "")).strip(),
                    risk_level=self._normalize_risk_level(item.get("risk_level")),
                    key_clause_values=key_values,
                    fields_to_compare=fields_to_compare,
                    source_chunk_index=task.chunk_index,
                )
            )
        return issues

    def _run_rule_checker(
        self,
        final_issues: list[FinalArbitratedIssue],
    ) -> tuple[list[CheckedFinding], list[CheckedFinding], list[CheckedFinding]]:
        """对仲裁后的结果执行增强版规则校验。"""
        accepted: list[CheckedFinding] = []
        needs_recheck: list[CheckedFinding] = []
        suppressed: list[CheckedFinding] = []
        retained_findings: list[CheckedFinding] = []

        for issue in final_issues:
            risk_level = self._normalize_risk_level(issue.risk_level)
            risk_score = RISK_SCORE_MAP.get(risk_level, 2) * 30
            duplicate_note = self._find_duplicate_reason(issue, retained_findings)
            if duplicate_note:
                suppressed_rule_hits = [
                    RuleHit(
                        rule_id="DEDUP-SEMANTIC-001",
                        priority="MAJOR",
                        result="SUPPRESS",
                        message=duplicate_note,
                    )
                ]
                suppressed.append(
                    CheckedFinding(
                        finding_id=issue.final_issue_id,
                        title=issue.title,
                        risk_level=risk_level,
                        risk_score=risk_score,
                        issue=issue.problem,
                        suggestion=issue.suggestion,
                        evidence=issue.evidence,
                        evidence_chunk_ids=issue.source_chunks,
                        dependency_ids=[],
                        source_chunk_ids=issue.source_chunks,
                        source_issue_ids=issue.source_issue_ids,
                        checker_status="suppressed",
                        checker_notes=[duplicate_note],
                        rule_hits=suppressed_rule_hits,
                    )
                )
                continue

            checker_notes: list[str] = []
            rule_hits: list[RuleHit] = []

            combined_text = " ".join(
                [issue.title, issue.problem, issue.suggestion, issue.evidence]
            )
            if any(token in combined_text for token in PLACEHOLDER_TOKENS):
                checker_notes.append("命中占位文本规则。")
                rule_hits.append(
                    RuleHit(
                        rule_id="QUALITY-PLACEHOLDER-001",
                        priority="BLOCKER",
                        result="FAIL",
                        message="输出包含占位文本。",
                    )
                )
            else:
                rule_hits.append(
                    RuleHit(
                        rule_id="QUALITY-PLACEHOLDER-001",
                        priority="BLOCKER",
                        result="PASS",
                        message="未命中占位文本。",
                    )
                )

            if risk_level == "高" and not issue.evidence.strip():
                checker_notes.append("高风险结论缺少证据摘要。")
                rule_hits.append(
                    RuleHit(
                        rule_id="RISK-HIGH-EVIDENCE-001",
                        priority="BLOCKER",
                        result="FAIL",
                        message="高风险结论必须提供 evidence。",
                    )
                )
            else:
                rule_hits.append(
                    RuleHit(
                        rule_id="RISK-HIGH-EVIDENCE-001",
                        priority="BLOCKER",
                        result="PASS",
                        message="高风险证据完整性通过。",
                    )
                )

            if risk_level in {"高", "中"} and not issue.source_chunks:
                checker_notes.append("中高风险结论缺少证据锚点。")
                rule_hits.append(
                    RuleHit(
                        rule_id="RISK-ANCHOR-001",
                        priority="BLOCKER",
                        result="FAIL",
                        message="中高风险结论必须提供 evidence_chunk_ids。",
                    )
                )
            else:
                rule_hits.append(
                    RuleHit(
                        rule_id="RISK-ANCHOR-001",
                        priority="BLOCKER",
                        result="PASS",
                        message="证据锚点字段存在。",
                    )
                )

            weak_evidence_note = self._get_weak_evidence_note(
                evidence=issue.evidence,
                risk_level=risk_level,
            )
            if weak_evidence_note:
                checker_notes.append(weak_evidence_note)
                rule_hits.append(
                    RuleHit(
                        rule_id="QUALITY-EVIDENCE-WEAK-001",
                        priority="MAJOR",
                        result="FAIL",
                        message=weak_evidence_note,
                    )
                )
            else:
                rule_hits.append(
                    RuleHit(
                        rule_id="QUALITY-EVIDENCE-WEAK-001",
                        priority="MAJOR",
                        result="PASS",
                        message="证据摘要强度通过。",
                    )
                )

            if not issue.suggestion.strip():
                checker_notes.append("修改建议为空，需人工复核。")
                rule_hits.append(
                    RuleHit(
                        rule_id="QUALITY-SUGGESTION-001",
                        priority="BLOCKER",
                        result="FAIL",
                        message="suggestion 不能为空。",
                    )
                )
            elif self._get_generic_suggestion_note(issue.suggestion):
                generic_suggestion_note = self._get_generic_suggestion_note(
                    issue.suggestion
                )
                checker_notes.append(generic_suggestion_note)
                rule_hits.append(
                    RuleHit(
                        rule_id="QUALITY-SUGGESTION-GENERIC-001",
                        priority="MAJOR",
                        result="FAIL",
                        message=generic_suggestion_note,
                    )
                )
            else:
                rule_hits.append(
                    RuleHit(
                        rule_id="QUALITY-SUGGESTION-001",
                        priority="BLOCKER",
                        result="PASS",
                        message="修改建议字段存在。",
                    )
                )
                rule_hits.append(
                    RuleHit(
                        rule_id="QUALITY-SUGGESTION-GENERIC-001",
                        priority="MAJOR",
                        result="PASS",
                        message="修改建议具备一定具体性。",
                    )
                )

            summary_title_note = self._get_summary_title_note(issue.title)
            if summary_title_note:
                checker_notes.append(summary_title_note)
                rule_hits.append(
                    RuleHit(
                        rule_id="QUALITY-SUMMARY-MIXED-001",
                        priority="MAJOR",
                        result="FAIL",
                        message=summary_title_note,
                    )
                )
            else:
                rule_hits.append(
                    RuleHit(
                        rule_id="QUALITY-SUMMARY-MIXED-001",
                        priority="MAJOR",
                        result="PASS",
                        message="标题未混入总结性内容。",
                    )
                )

            checked_finding = CheckedFinding(
                finding_id=issue.final_issue_id,
                title=issue.title,
                risk_level=risk_level,
                risk_score=risk_score,
                issue=issue.problem,
                suggestion=issue.suggestion,
                evidence=issue.evidence,
                evidence_chunk_ids=issue.source_chunks,
                dependency_ids=[],
                source_chunk_ids=issue.source_chunks,
                source_issue_ids=issue.source_issue_ids,
                checker_status="accepted" if not checker_notes else "needs_recheck",
                checker_notes=checker_notes,
                rule_hits=rule_hits,
            )
            retained_findings.append(checked_finding)
            if checked_finding.checker_status == "accepted":
                accepted.append(checked_finding)
            else:
                needs_recheck.append(checked_finding)

        return accepted, needs_recheck, suppressed

    def _precheck_consolidate_final_issues(
        self,
        final_issues: list[FinalArbitratedIssue],
    ) -> tuple[list[FinalArbitratedIssue], int]:
        """进入 checker 前再做一轮主题级收束。"""
        consolidated: list[FinalArbitratedIssue] = []
        merge_count = 0

        for issue in final_issues:
            target_index = self._find_precheck_merge_target(issue, consolidated)
            if target_index is None:
                consolidated.append(issue)
                continue
            consolidated[target_index] = self._merge_final_issues(
                consolidated[target_index],
                issue,
            )
            merge_count += 1

        consolidated.sort(
            key=lambda item: (
                -RISK_SCORE_MAP.get(item.risk_level, 0),
                min(item.source_chunks) if item.source_chunks else 0,
                item.final_issue_id,
            )
        )
        return consolidated, merge_count

    def _find_precheck_merge_target(
        self,
        issue: FinalArbitratedIssue,
        consolidated: list[FinalArbitratedIssue],
    ) -> int | None:
        """寻找可归并的已有 final issue。"""
        current_title = self._normalize_free_text(issue.title)
        current_title_core = self._canonicalize_title(issue.title)
        current_problem = self._normalize_free_text(issue.problem)
        current_family = self._infer_topic_family(
            " ".join([issue.title, issue.problem, issue.suggestion])
        )
        current_sources = {str(item) for item in issue.source_chunks}

        for index, retained in enumerate(consolidated):
            retained_title = self._normalize_free_text(retained.title)
            retained_title_core = self._canonicalize_title(retained.title)
            retained_problem = self._normalize_free_text(retained.problem)
            retained_family = self._infer_topic_family(
                " ".join([retained.title, retained.problem, retained.suggestion])
            )
            retained_sources = {str(item) for item in retained.source_chunks}
            source_overlap = bool(current_sources.intersection(retained_sources))
            title_similarity = max(
                self._similarity(current_title, retained_title),
                self._similarity(current_title_core, retained_title_core),
            )
            problem_similarity = self._similarity(current_problem, retained_problem)

            if (
                current_title_core
                and retained_title_core
                and current_title_core == retained_title_core
                and (problem_similarity >= 0.58 or source_overlap)
            ):
                return index

            if (
                current_family
                and retained_family
                and current_family == retained_family
                and (
                    title_similarity >= 0.68
                    or problem_similarity >= 0.76
                    or (source_overlap and title_similarity >= 0.52)
                )
            ):
                return index

            if source_overlap and title_similarity >= 0.82:
                return index

        return None

    def _merge_final_issues(
        self,
        retained: FinalArbitratedIssue,
        incoming: FinalArbitratedIssue,
    ) -> FinalArbitratedIssue:
        """合并主题相近的 final issue。"""
        merged_title = self._pick_better_final_issue_title(
            retained.title,
            incoming.title,
        )
        merged_problem = self._merge_text_blocks([retained.problem, incoming.problem])
        merged_suggestion = self._pick_better_text_block(
            retained.suggestion,
            incoming.suggestion,
        )
        merged_evidence = self._merge_text_blocks([retained.evidence, incoming.evidence])
        merged_chunks = sorted({*retained.source_chunks, *incoming.source_chunks})
        merged_issue_ids = self._unique_texts(
            retained.source_issue_ids + incoming.source_issue_ids
        )
        merged_risk_level = self._pick_higher_risk_level(
            retained.risk_level,
            incoming.risk_level,
        )
        return FinalArbitratedIssue(
            final_issue_id=retained.final_issue_id,
            title=merged_title,
            risk_level=merged_risk_level,
            problem=merged_problem,
            suggestion=merged_suggestion,
            evidence=merged_evidence,
            source_chunks=merged_chunks,
            source_issue_ids=merged_issue_ids,
        )

    def _find_duplicate_reason(
        self,
        issue: FinalArbitratedIssue,
        retained_findings: list[CheckedFinding],
    ) -> str:
        """识别近义标题和同证据重复。"""
        current_title = self._normalize_free_text(issue.title)
        current_title_core = self._canonicalize_title(issue.title)
        current_issue = self._normalize_free_text(issue.problem)
        current_evidence = self._normalize_free_text(issue.evidence)
        current_sources = {str(item) for item in issue.source_chunks}

        for retained in retained_findings:
            retained_title = self._normalize_free_text(retained.title)
            retained_title_core = self._canonicalize_title(retained.title)
            retained_issue = self._normalize_free_text(retained.issue)
            retained_evidence = self._normalize_free_text(retained.evidence)
            retained_sources = {str(item) for item in retained.source_chunk_ids}
            source_overlap = bool(current_sources.intersection(retained_sources))
            title_similarity = max(
                self._similarity(current_title, retained_title),
                self._similarity(current_title_core, retained_title_core),
            )
            issue_similarity = self._similarity(current_issue, retained_issue)
            evidence_similarity = self._similarity(
                current_evidence,
                retained_evidence,
            )

            if self._is_same_evidence_duplicate(
                current_evidence=current_evidence,
                retained_evidence=retained_evidence,
                source_overlap=source_overlap,
                title_similarity=title_similarity,
            ):
                return f"与既有 finding “{retained.title}”证据高度重复，已进入 suppressed。"

            if (
                current_title_core
                and retained_title_core
                and current_title_core == retained_title_core
                and (
                    evidence_similarity >= 0.72
                    or issue_similarity >= 0.82
                    or (source_overlap and issue_similarity >= 0.68)
                )
            ):
                return f"与既有 finding “{retained.title}”属于近义标题重复，已进入 suppressed。"

            if title_similarity >= 0.9 and issue_similarity >= 0.8:
                return f"与既有 finding “{retained.title}”语义高度相近，已进入 suppressed。"

        return ""

    def _is_same_evidence_duplicate(
        self,
        *,
        current_evidence: str,
        retained_evidence: str,
        source_overlap: bool,
        title_similarity: float,
    ) -> bool:
        """判断是否属于同证据重复。"""
        if not current_evidence or not retained_evidence:
            return False

        shorter, longer = sorted(
            [current_evidence, retained_evidence],
            key=len,
        )
        if len(shorter) >= 24 and shorter in longer and (source_overlap or title_similarity >= 0.55):
            return True

        evidence_similarity = self._similarity(current_evidence, retained_evidence)
        if evidence_similarity >= 0.92 and (source_overlap or title_similarity >= 0.6):
            return True
        return False

    def _get_weak_evidence_note(
        self,
        *,
        evidence: str,
        risk_level: str,
    ) -> str:
        """识别弱证据。"""
        normalized_evidence = self._normalize_free_text(evidence)
        if risk_level in {"高", "中"} and len(normalized_evidence) < 18:
            return "中高风险结论的证据摘要过短，需人工复核。"
        if any(token in evidence for token in GENERIC_EVIDENCE_TOKENS):
            return "证据摘要过于模板化，缺少具体原文支撑。"
        return ""

    def _get_generic_suggestion_note(self, suggestion: str) -> str:
        """识别空泛建议。"""
        normalized_suggestion = self._normalize_free_text(suggestion)
        if len(normalized_suggestion) < 16:
            return "修改建议过短，缺少可执行细节。"
        if any(token in suggestion for token in GENERIC_SUGGESTION_TOKENS):
            return "修改建议过于空泛，缺少具体改写方案。"
        return ""

    def _get_summary_title_note(self, title: str) -> str:
        """识别总结性标题混入 finding。"""
        if "**" in title:
            return "标题混入 markdown 总结样式，需人工复核。"
        if any(token in title for token in SUMMARY_TITLE_TOKENS) and len(title.strip()) >= 8:
            return "标题更像总结性段落，而非单条 finding。"
        return ""

    def _save_run_result(
        self,
        *,
        snapshot: EvidenceSnapshot,
        tasks: list[SpecialistTask],
        chunk_opinions: list[ChunkAnalysisOpinion],
        conflicts: list[ConflictRecord],
        decisions: list[ArbitrationDecision],
        precheck_merge_count: int,
        accepted_findings: list[CheckedFinding],
        needs_recheck_findings: list[CheckedFinding],
        suppressed_findings: list[CheckedFinding],
        elapsed_seconds: float,
        stage_timings: list[StageTiming],
    ) -> dict[str, str]:
        """保存文本与 JSON 结果。"""
        result_dir = ensure_multi_agent_demo_result_dir(self.demo_config)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_stem = Path(snapshot.file_path).stem
        text_path = (
            result_dir / f"scheme_a_merge_rule_engine_{file_stem}_{timestamp}.txt"
        )
        json_path = (
            result_dir / f"scheme_a_merge_rule_engine_{file_stem}_{timestamp}.json"
        )

        task_counter = Counter(task.specialist_name for task in tasks)
        chunk_task_counter = Counter(task.chunk_id for task in tasks)
        rule_counter = Counter(
            hit.rule_id
            for item in accepted_findings + needs_recheck_findings + suppressed_findings
            for hit in item.rule_hits
            if hit.result != "PASS"
        )
        zero_task_chunk_count = len(snapshot.chunks) - len(chunk_task_counter)
        multi_task_chunk_count = sum(1 for count in chunk_task_counter.values() if count >= 2)

        lines: list[str] = [
            "方案 A：归并仲裁 + 规则引擎 Demo 运行结果",
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"源文件路径：{snapshot.file_path}",
            f"源文件类型：{snapshot.file_type}",
            f"证据快照版本：{snapshot.snapshot_id}",
            f"分块数量：{len(snapshot.chunks)}",
            f"任务数量：{len(tasks)}",
            f"未派发任务分块数：{zero_task_chunk_count}",
            f"双 specialist 分块数：{multi_task_chunk_count}",
            f"结构化意见池数量：{sum(len(item.issues) for item in chunk_opinions)}",
            f"冲突数量：{len(conflicts)}",
            f"预归并收束数量：{precheck_merge_count}",
            f"accepted 数量：{len(accepted_findings)}",
            f"needs_recheck 数量：{len(needs_recheck_findings)}",
            f"suppressed 数量：{len(suppressed_findings)}",
            f"总计耗时（秒）：{elapsed_seconds:.3f}",
            "",
            "## 阶段耗时",
        ]
        for item in stage_timings:
            lines.append(f"- {item.stage_name}：{item.elapsed_seconds:.3f} 秒")

        lines.extend(
            [
                "",
            "## Specialist 任务分布",
            ]
        )
        for specialist_name, count in sorted(task_counter.items()):
            lines.append(f"- {specialist_name}：{count}")

        lines.append("")
        lines.append("## 规则命中统计")
        if not rule_counter:
            lines.append("未命中失败规则。")
        else:
            for rule_id, count in sorted(rule_counter.items()):
                lines.append(f"- {rule_id}：{count}")

        lines.append("")
        lines.append("## 仲裁决策")
        if not decisions:
            lines.append("未触发额外仲裁决策。")
        else:
            for decision in decisions:
                lines.append(f"- 决策标题：{decision.title}")
                lines.append(f"  动作：{decision.action}")
                lines.append(f"  理由：{decision.rationale}")

        self._append_checked_findings(lines, "Accepted Findings", accepted_findings)
        self._append_checked_findings(
            lines, "Needs Recheck Findings", needs_recheck_findings
        )
        self._append_checked_findings(lines, "Suppressed Findings", suppressed_findings)

        lines.append("## 审计摘要")
        lines.append(f"- 证据快照：{snapshot.snapshot_id}")
        lines.append(f"- 任务数：{len(tasks)}")
        lines.append(f"- accepted：{len(accepted_findings)}")
        lines.append(f"- needs_recheck：{len(needs_recheck_findings)}")
        lines.append(f"- suppressed：{len(suppressed_findings)}")
        lines.append(f"- 总耗时（秒）：{elapsed_seconds:.3f}")

        with open(text_path, "w", encoding="utf-8") as file:
            file.write("\n".join(lines).strip() + "\n")

        payload = {
            "snapshot": asdict(snapshot),
            "tasks": [asdict(task) for task in tasks],
            "chunk_opinions": [asdict(item) for item in chunk_opinions],
            "conflicts": [asdict(item) for item in conflicts],
            "decisions": [asdict(item) for item in decisions],
            "precheck_merge_count": precheck_merge_count,
            "accepted_findings": [self._checked_finding_to_dict(item) for item in accepted_findings],
            "needs_recheck_findings": [
                self._checked_finding_to_dict(item)
                for item in needs_recheck_findings
            ],
            "suppressed_findings": [
                self._checked_finding_to_dict(item) for item in suppressed_findings
            ],
            "stage_timings": [asdict(item) for item in stage_timings],
            "elapsed_seconds": elapsed_seconds,
        }
        with open(json_path, "w", encoding="utf-8") as file:
            json.dump(payload, file, ensure_ascii=False, indent=2)

        return {"text": str(text_path), "json": str(json_path)}

    def _append_checked_findings(
        self,
        lines: list[str],
        title: str,
        findings: list[CheckedFinding],
    ) -> None:
        """将校验后的结果写入文本报告。"""
        lines.append(f"## {title}")
        if not findings:
            lines.append("无")
            lines.append("")
            return
        for item in findings:
            lines.append(f"### {item.title}")
            lines.append(f"状态：{item.checker_status}")
            lines.append(f"风险等级：{item.risk_level}")
            lines.append(f"风险分值：{item.risk_score}")
            lines.append(f"来源分块：{','.join(map(str, item.source_chunk_ids)) or '无'}")
            lines.append(f"问题说明：{item.issue or '无'}")
            lines.append(f"修改建议：{item.suggestion or '无'}")
            lines.append(f"证据摘要：{self._truncate_text(item.evidence, limit=180) or '无'}")
            if item.checker_notes:
                lines.append(f"checker 说明：{'；'.join(item.checker_notes)}")
            if item.rule_hits:
                lines.append(
                    "规则命中："
                    + "；".join(
                        f"{hit.rule_id}:{hit.result}" for hit in item.rule_hits
                    )
                )
            lines.append("")

    def _checked_finding_to_dict(self, item: CheckedFinding) -> dict[str, Any]:
        """转换 CheckedFinding 为可序列化对象。"""
        payload = asdict(item)
        payload["rule_hits"] = [asdict(hit) for hit in item.rule_hits]
        return payload

    def _extract_chunk_compare_metadata(
        self,
        text: str,
    ) -> tuple[dict[str, str], list[str]]:
        """从分块中提取最小比对字段。"""
        key_values: dict[str, str] = {}
        compare_fields: list[str] = []

        payment_match = re.search(
            r"([0-9一二三四五六七八九十]+个?工作日内支付|[0-9一二三四五六七八九十]+日内支付)",
            text,
        )
        if payment_match:
            key_values["付款期限"] = payment_match.group(1)
            compare_fields.append("付款期限")

        penalty_match = re.search(
            r"违约金[^。；\n]{0,20}([0-9]+(?:\.[0-9]+)?‰|[0-9]+(?:\.[0-9]+)?%)",
            text,
        )
        if penalty_match:
            key_values["违约金比例"] = penalty_match.group(1)
            compare_fields.append("违约金比例")

        dispute_match = re.search(
            r"(项目所在地人民法院|[^\n。；]{2,20}仲裁委员会|[^\n。；]{2,20}人民法院)",
            text,
        )
        if dispute_match and ("仲裁" in dispute_match.group(1) or "法院" in dispute_match.group(1)):
            key_values["争议解决地"] = dispute_match.group(1)
            compare_fields.append("争议解决地")

        defect_match = re.search(
            r"缺陷责任期[^。；\n]{0,20}([0-9一二三四五六七八九十]+(?:个月|年|日))",
            text,
        )
        if defect_match:
            key_values["缺陷责任期"] = defect_match.group(1)
            compare_fields.append("缺陷责任期")

        return key_values, self._unique_texts(compare_fields)

    def _infer_chunk_tags(self, text: str) -> list[str]:
        """根据关键词抽取轻量标签。"""
        tags: list[str] = []
        for profile in SPECIALIST_PROFILES:
            tags.extend(keyword for keyword in profile.keywords if keyword in text)
        return self._unique_texts(tags)

    def _extract_reference_tokens(self, text: str) -> list[str]:
        """抽取引用标记。"""
        return self._unique_texts(REFERENCE_PATTERN.findall(text))

    def _extract_definition_pairs(self, text: str) -> list[tuple[str, str]]:
        """抽取轻量定义项。"""
        pairs: list[tuple[str, str]] = []
        for match in re.finditer(
            r"(?:“(?P<quoted>[^”]{2,24})”|(?P<plain>[\u4e00-\u9fa5A-Za-z]{2,24}))\s*(?:系指|是指|指)\s*(?P<desc>[^。；\n]{4,120})",
            text,
        ):
            term = (match.group("quoted") or match.group("plain") or "").strip()
            description = (match.group("desc") or "").strip()
            if term and description:
                pairs.append((term, description))
        return pairs

    def _normalize_key(self, value: str) -> str:
        """标准化分组键。"""
        return re.sub(r"\s+", "", str(value).strip()).lower()

    def _contains_placeholder_signal(self, text: str) -> bool:
        """识别空白、占位、待补充等异常信号。"""
        return any(
            token in text
            for token in (
                "留空",
                "待定",
                "待补充",
                "另行约定",
                "另行协商",
                "详见附件",
                "填写",
            )
        )

    def _normalize_free_text(self, value: str) -> str:
        """标准化自由文本。"""
        return re.sub(r"\s+", "", str(value).strip())

    def _canonicalize_title(self, title: str) -> str:
        """提取标题核心，用于近义去重。"""
        normalized_title = self._normalize_free_text(title)
        for prefix in (
            "同一事项建议冲突：",
            "关键字段冲突：",
            "跨块联动冲突：",
            "漏判风险冲突：",
        ):
            normalized_title = normalized_title.replace(
                self._normalize_free_text(prefix),
                "",
            )
        for token in TITLE_CANONICALIZE_TOKENS:
            normalized_title = normalized_title.replace(token, "")
        return normalized_title

    def _infer_topic_family(self, text: str) -> str:
        """识别主题族，用于进入 checker 前的收束。"""
        family_rules = [
            ("payment", ("付款", "支付", "结算", "价款", "预付款", "进度款")),
            ("liability", ("责任", "违约", "赔偿", "违约金", "免责", "索赔")),
            ("termination", ("解除", "终止", "退场", "撤场")),
            ("dispute", ("争议", "仲裁", "法院", "诉讼", "管辖")),
            ("definition", ("定义", "附件", "附表", "系指", "是指")),
            ("acceptance", ("验收", "交付", "试运行", "缺陷责任期")),
            ("procedure", ("通知", "程序", "期限", "审批", "确认")),
            ("authority", ("主体", "授权", "签署", "盖章")),
        ]
        for family_name, tokens in family_rules:
            if any(token in text for token in tokens):
                return family_name
        return ""

    def _pick_better_final_issue_title(self, left: str, right: str) -> str:
        """挑选更适合作为聚合标题的文本。"""
        left_title = left.strip()
        right_title = right.strip()
        if not left_title:
            return right_title
        if not right_title:
            return left_title

        left_penalty = self._title_penalty(left_title)
        right_penalty = self._title_penalty(right_title)
        if left_penalty != right_penalty:
            return left_title if left_penalty < right_penalty else right_title

        if abs(len(left_title) - len(right_title)) >= 10:
            return left_title if len(left_title) < len(right_title) else right_title
        return left_title

    def _title_penalty(self, title: str) -> int:
        """为偏泛化标题增加惩罚分。"""
        penalty = 0
        if "同一事项建议冲突：" in title:
            penalty += 2
        if "跨块联动冲突：" in title or "漏判风险冲突：" in title:
            penalty += 3
        if any(token in title for token in SUMMARY_TITLE_TOKENS):
            penalty += 3
        return penalty

    def _merge_text_blocks(self, texts: list[str]) -> str:
        """合并多段文本，同时做顺序去重。"""
        merged_parts: list[str] = []
        for text in texts:
            clean_text = str(text).strip()
            normalized_text = self._normalize_free_text(clean_text)
            if not clean_text:
                continue
            if any(
                normalized_text and normalized_text in self._normalize_free_text(item)
                for item in merged_parts
            ):
                continue
            merged_parts.append(clean_text)
        return "；".join(merged_parts)

    def _pick_better_text_block(self, left: str, right: str) -> str:
        """挑选更具体的建议文本。"""
        left_text = left.strip()
        right_text = right.strip()
        if not left_text:
            return right_text
        if not right_text:
            return left_text

        left_note = self._get_generic_suggestion_note(left_text)
        right_note = self._get_generic_suggestion_note(right_text)
        if bool(left_note) != bool(right_note):
            return right_text if left_note else left_text
        return left_text if len(left_text) >= len(right_text) else right_text

    def _pick_higher_risk_level(self, left: str, right: str) -> str:
        """返回更高的风险等级。"""
        left_score = RISK_SCORE_MAP.get(self._normalize_risk_level(left), 0)
        right_score = RISK_SCORE_MAP.get(self._normalize_risk_level(right), 0)
        return self._normalize_risk_level(left if left_score >= right_score else right)

    def _similarity(self, left: str, right: str) -> float:
        """计算相似度。"""
        if not left or not right:
            return 0.0
        return SequenceMatcher(None, left, right).ratio()

    def _truncate_text(self, text: str, *, limit: int = 120) -> str:
        """截断长文本。"""
        clean_text = re.sub(r"\s+", " ", text.strip())
        if len(clean_text) <= limit:
            return clean_text
        return clean_text[: limit - 3] + "..."

    def _unique_texts(self, items: list[str]) -> list[str]:
        """保持顺序去重。"""
        result: list[str] = []
        for item in items:
            text = str(item).strip()
            if text and text not in result:
                result.append(text)
        return result


class _FakeChatCompletions:
    """用于本地结构验证的假模型。"""

    def create(self, *, model: str, messages: list, tools: list, **kwargs):
        user_prompt = messages[-1]["content"]
        if "付款与结算 specialist" in user_prompt and "30个工作日内支付全部合同价款" in user_prompt:
            content = """
【修改点1】付款期限偏长
【原文】甲方应在验收合格后30个工作日内支付全部合同价款。
【风险分析】付款期限较长，容易拉长承包人回款周期并引发履约争议。
【风险等级】高
【修改后的内容】甲方应在验收合格并收到合规发票后10个工作日内支付全部合同价款。
【修改理由】缩短付款期限并补足付款触发条件。
【风险类型】付款期限
""".strip()
        elif "责任与违约 specialist" in user_prompt and "违约金比例留空" in user_prompt:
            content = """
【修改点1】违约金比例留空
【原文】违约责任：任何一方违约均应承担损失，违约金比例留空。
【风险分析】违约金比例为空白，导致责任标准无法执行。
【风险等级】高
【修改后的内容】违约责任：任何一方违约均应承担损失，违约金比例建议明确为合同金额的5%。
【修改理由】补足关键责任字段，避免履约争议。
【风险类型】违约金比例
""".strip()
        elif "终止与争议 specialist" in user_prompt and "项目所在地法院" in user_prompt:
            content = """
【修改点1】争议解决条款不完整
【原文】争议提交项目所在地法院。
【风险分析】争议解决方式仅写法院，未明确前置协商或具体管辖层级。
【风险等级】中
【修改后的内容】争议应先协商解决，协商不成的，提交项目所在地有管辖权的人民法院诉讼解决。
【修改理由】补足争议解决程序，降低管辖争议。
【风险类型】争议解决
""".strip()
        else:
            content = ""

        return SimpleNamespace(
            choices=[SimpleNamespace(message=SimpleNamespace(content=content))]
        )


class _FakeLLM:
    """用于本地结构验证的假客户端。"""

    def __init__(self):
        self.chat = SimpleNamespace(completions=_FakeChatCompletions())


async def _main_test_merge_rule_engine_demo():
    """本地最小结构验证。"""
    with tempfile.TemporaryDirectory(prefix="scheme_a_merge_rule_engine_test_") as temp_dir:
        sample_path = Path(temp_dir) / "sample_contract.docx"
        result_dir = Path(temp_dir) / "result"
        document = Document()
        document.add_paragraph(
            "定义条款：本合同所称“验收合格”是指甲方书面确认工作成果符合约定标准。"
        )
        document.add_paragraph(
            "付款条款：甲方应在验收合格后30个工作日内支付全部合同价款。"
        )
        document.add_paragraph(
            "违约责任：任何一方违约均应承担损失，违约金比例留空。争议提交项目所在地法院。"
        )
        document.save(sample_path)

        demo_config = MultiAgentDemoConfig(
            model=MultiAgentDemoModelConfig(
                model_name="fake-scheme-a-model",
                api_key="fake-key",
                api_base="https://example.com/v1",
                temperature=0.0,
                top_p=1.0,
                max_tokens=1024,
            ),
            chunk_size=40,
            max_concurrent_reviews=2,
            result_dir=str(result_dir),
        )
        service = MergeRuleEngineReviewDemo(
            config=demo_config,
            llm_client=_FakeLLM(),
        )
        result = await service.run(str(sample_path))
        print("scheme_a_merge_rule_engine_demo self test result:")
        print(
            {
                "chunk_count": result["chunk_count"],
                "task_count": result["task_count"],
                "accepted_count": result["accepted_count"],
                "needs_recheck_count": result["needs_recheck_count"],
                "suppressed_count": result["suppressed_count"],
                "elapsed_seconds": round(result["elapsed_seconds"], 3),
                "result_file_exists": Path(result["result_file_path"]).exists(),
                "result_json_exists": Path(result["result_json_path"]).exists(),
            }
        )
        assert result["chunk_count"] >= 3
        assert result["task_count"] >= 3
        assert result["accepted_count"] >= 1
        assert result["elapsed_seconds"] >= 0
        assert Path(result["result_file_path"]).exists()
        assert Path(result["result_json_path"]).exists()
        print("scheme_a_merge_rule_engine_demo self test passed")


async def _main_run_real_model_demo(file_path: str):
    """使用真实模型运行 demo。"""
    service = MergeRuleEngineReviewDemo(config=get_multi_agent_demo_config())
    result = await service.run(file_path)
    print("scheme_a_merge_rule_engine_demo real run result:")
    print(
        {
            "chunk_count": result["chunk_count"],
            "task_count": result["task_count"],
            "accepted_count": result["accepted_count"],
            "needs_recheck_count": result["needs_recheck_count"],
            "suppressed_count": result["suppressed_count"],
            "elapsed_seconds": round(result["elapsed_seconds"], 3),
            "result_file_path": result["result_file_path"],
            "result_json_path": result["result_json_path"],
        }
    )


def _build_arg_parser() -> argparse.ArgumentParser:
    """构造命令行参数。"""
    parser = argparse.ArgumentParser(description="方案 A：归并仲裁 + 规则引擎 Demo")
    parser.add_argument(
        "--file",
        dest="file_path",
        help="真实模型运行时的合同文件路径（doc/docx）",
    )
    return parser


if __name__ == "__main__":
    arguments = _build_arg_parser().parse_args()
    if arguments.file_path:
        asyncio.run(_main_run_real_model_demo(arguments.file_path))
    else:
        asyncio.run(_main_test_merge_rule_engine_demo())
