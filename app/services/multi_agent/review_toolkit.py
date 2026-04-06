"""multi_agent 通用审阅工具。"""

from __future__ import annotations

import asyncio
import logging
import tempfile
from dataclasses import dataclass
from pathlib import Path
from types import SimpleNamespace
from typing import Any

from docx import Document

from app.services.contract_review import ContractReviewService
from app.utils.content_slicer import split_text_by_length
from app.utils.contract_parser import ContractParser
from app.utils.document_parsing import doc2docx
from prompts.llm_prompt_vars import (
    REVIEW_SYSTEM_PROMPT,
    build_contract_review_prompt,
)

try:
    from .config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        get_multi_agent_demo_config,
        init_multi_agent_demo_llm,
    )
except ImportError:
    from app.services.multi_agent.config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        get_multi_agent_demo_config,
        init_multi_agent_demo_llm,
    )

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class ParsedContract:
    """解析后的合同文本。"""

    file_path: str
    file_type: str
    text: str


class MultiAgentReviewToolkit(ContractReviewService):
    """为多 agent demo 提供解析、分块和审阅基础能力。"""

    def __init__(
        self,
        config: MultiAgentDemoConfig | None = None,
        llm_client: Any | None = None,
    ):
        self.demo_config = config or get_multi_agent_demo_config()
        self.model_config = self.demo_config.model
        self.llm = llm_client or init_multi_agent_demo_llm(self.model_config)
        self.mcp_client = None
        self.rag_service = None

    async def parse_contract_file(self, file_path: str) -> ParsedContract:
        """解析本地 doc/docx 文件并返回纯文本。"""
        source_path = Path(file_path).expanduser().resolve()
        if not source_path.exists():
            raise FileNotFoundError(f"待审阅文件不存在: {source_path}")

        suffix = source_path.suffix.lower()
        if suffix not in {".doc", ".docx"}:
            raise ValueError("demo 仅支持 doc/docx 文件")

        if suffix == ".docx":
            text = await ContractParser.read_docx_text(str(source_path))
            return ParsedContract(
                file_path=str(source_path),
                file_type="docx",
                text=text,
            )

        with tempfile.TemporaryDirectory(prefix="multi_agent_doc_") as temp_dir:
            converted_path = Path(temp_dir) / f"{source_path.stem}.docx"
            doc2docx(str(source_path), str(converted_path))
            if not converted_path.exists():
                raise RuntimeError(f"DOC 转 DOCX 失败: {source_path}")
            text = await ContractParser.read_docx_text(str(converted_path))

        return ParsedContract(
            file_path=str(source_path),
            file_type="doc",
            text=text,
        )

    def split_contract_text(self, contract_text: str) -> list[str]:
        """按主链路规则拆分合同文本。"""
        chunks = split_text_by_length(
            contract_text,
            max_length=self.demo_config.chunk_size,
        )
        return [chunk for chunk in chunks if chunk.strip()]

    async def review_chunk(
        self,
        *,
        chunk_text: str,
        stance: str,
        intensity: str,
        contract_type: str,
        context: str = "",
    ) -> list[dict[str, Any]]:
        """审阅单个文本分块。"""
        try:
            base_prompt, contract_type_prompt = self._load_review_prompts(contract_type)
            review_prompt = build_contract_review_prompt(
                base_prompt=base_prompt,
                contract_type_prompt=contract_type_prompt,
                stance=stance,
                intensity=intensity,
                contract_type=contract_type,
                context=context,
                chunk_text=chunk_text,
            )
            messages = [
                {"role": "system", "content": REVIEW_SYSTEM_PROMPT},
                {"role": "user", "content": review_prompt},
            ]
            response = self.llm.chat.completions.create(
                model=self.model_config.model_name,
                messages=messages,
                tools=[],
                temperature=self.model_config.temperature,
                top_p=self.model_config.top_p,
                max_tokens=self.model_config.max_tokens,
            )
            review_result = response.choices[0].message.content or ""
            return self._parse_review_result(review_result)
        except Exception as error:
            logger.warning("分块审阅失败，返回空结果: %s", error)
            return []

    def _load_review_prompts(self, contract_type: str) -> tuple[str, str]:
        """加载与主链路一致的审阅 prompt 模板。"""
        base_prompt_dir = Path(__file__).resolve().parents[3] / "prompts"
        base_prompt_path = base_prompt_dir / "contract_reviewer_prompt_unified.txt"

        if contract_type == "基建类合同":
            contract_type_prompt_file = "contract_reviewer_prompt_build.txt"
        elif contract_type == "货物类合同":
            contract_type_prompt_file = "contract_reviewer_prompt_sales.txt"
        elif contract_type == "服务类合同":
            contract_type_prompt_file = "contract_reviewer_prompt_service.txt"
        else:
            contract_type_prompt_file = "contract_reviewer_prompt_base.txt"

        contract_type_prompt_path = base_prompt_dir / contract_type_prompt_file
        with open(base_prompt_path, "r", encoding="utf-8") as file:
            base_prompt = file.read()
        with open(contract_type_prompt_path, "r", encoding="utf-8") as file:
            contract_type_prompt = file.read()
        return base_prompt, contract_type_prompt


class _FakeChatCompletions:
    """用于本地结构验证的假模型。"""

    def create(self, *, model: str, messages: list, tools: list, **kwargs):
        fake_review_result = """
【修改点1】付款条款表述不清
【原文】乙方完成工作后甲方付款。
【风险分析】付款触发条件和付款时间不明确，容易引发履约争议。
【风险等级】中
【修改后的内容】乙方完成工作并经甲方书面验收通过后，甲方应于10个工作日内支付合同款项。
【修改理由】补足验收条件和付款期限，减少争议。
【风险类型】付款条款
""".strip()
        return SimpleNamespace(
            choices=[
                SimpleNamespace(
                    message=SimpleNamespace(content=fake_review_result),
                )
            ]
        )


class _FakeLLM:
    """用于本地结构验证的假客户端。"""

    def __init__(self):
        self.chat = SimpleNamespace(completions=_FakeChatCompletions())


async def _main_test_review_toolkit():
    """本地最小结构验证。"""
    with tempfile.TemporaryDirectory(prefix="multi_agent_demo_test_") as temp_dir:
        sample_path = Path(temp_dir) / "sample_contract.docx"
        document = Document()
        document.add_paragraph("乙方完成工作后甲方付款。")
        document.add_paragraph("任何一方违约应承担责任。")
        document.save(sample_path)

        demo_config = MultiAgentDemoConfig(
            model=MultiAgentDemoModelConfig(
                model_name="fake-review-toolkit-model",
                api_key="fake-key",
                api_base="https://example.com/v1",
                temperature=0.0,
                top_p=1.0,
                max_tokens=512,
            ),
            chunk_size=1000,
        )
        service = MultiAgentReviewToolkit(
            config=demo_config,
            llm_client=_FakeLLM(),
        )

        parsed_contract = await service.parse_contract_file(str(sample_path))
        chunks = service.split_contract_text(parsed_contract.text)
        result = await service.review_chunk(
            chunk_text=chunks[0],
            stance="甲方",
            intensity="标准",
            contract_type="通用",
            context="这是第 1 个分块，共 1 个。",
        )
        print("review_toolkit self test result:")
        print(
            {
                "file_type": parsed_contract.file_type,
                "chunk_count": len(chunks),
                "modification_count": len(result),
                "first_risk_level": result[0]["risk_level"],
            }
        )
        assert parsed_contract.file_type == "docx"
        assert len(chunks) == 1
        assert len(result) == 1
        assert result[0]["risk_level"] == "中"
        print("review_toolkit self test passed")


if __name__ == "__main__":
    asyncio.run(_main_test_review_toolkit())
