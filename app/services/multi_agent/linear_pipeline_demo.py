"""线性顺序流水线多 agent demo。"""

from __future__ import annotations

import asyncio
import tempfile
import time
from datetime import datetime
from pathlib import Path
from types import SimpleNamespace
from typing import Any

from docx import Document

try:
    from .config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        ensure_multi_agent_demo_result_dir,
    )
    from .review_toolkit import MultiAgentReviewToolkit
except ImportError:
    from app.services.multi_agent.config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        ensure_multi_agent_demo_result_dir,
    )
    from app.services.multi_agent.review_toolkit import MultiAgentReviewToolkit


class LinearPipelineReviewDemo(MultiAgentReviewToolkit):
    """按解析、分块、异步审阅、结果落盘顺序执行的 demo。"""

    async def run(
        self,
        file_path: str,
        *,
        stance: str | None = None,
        intensity: str | None = None,
        contract_type: str | None = None,
    ) -> dict[str, Any]:
        """执行线性顺序流水线 demo。"""
        start_time = time.perf_counter()
        parsed_contract = await self.parse_contract_file(file_path)
        chunks = self.split_contract_text(parsed_contract.text)
        chunk_results = await self._review_all_chunks(
            chunks=chunks,
            stance=stance or self.demo_config.default_stance,
            intensity=intensity or self.demo_config.default_intensity,
            contract_type=contract_type or self.demo_config.default_contract_type,
        )
        elapsed_seconds = time.perf_counter() - start_time
        modification_count = sum(
            len(item["modifications"]) for item in chunk_results
        )
        result_file_path = self._save_run_result(
            parsed_contract=parsed_contract,
            chunk_results=chunk_results,
            elapsed_seconds=elapsed_seconds,
        )
        return {
            "demo_type": "linear_pipeline",
            "file_path": parsed_contract.file_path,
            "file_type": parsed_contract.file_type,
            "chunk_count": len(chunks),
            "modification_count": modification_count,
            "elapsed_seconds": elapsed_seconds,
            "result_file_path": result_file_path,
            "chunk_results": chunk_results,
        }

    async def _review_all_chunks(
        self,
        *,
        chunks: list[str],
        stance: str,
        intensity: str,
        contract_type: str,
    ) -> list[dict[str, Any]]:
        """异步审阅全部分块并保持结果顺序。"""
        semaphore = asyncio.Semaphore(self.demo_config.max_concurrent_reviews)
        ordered_results: list[dict[str, Any] | None] = [None] * len(chunks)

        async def _review_single(index: int, chunk_text: str):
            async with semaphore:
                context = f"这是第 {index + 1} 个分块，共 {len(chunks)} 个。"
                modifications = await self.review_chunk(
                    chunk_text=chunk_text,
                    stance=stance,
                    intensity=intensity,
                    contract_type=contract_type,
                    context=context,
                )
                ordered_results[index] = {
                    "chunk_index": index + 1,
                    "chunk_text": chunk_text,
                    "modifications": modifications,
                }

        await asyncio.gather(
            *[_review_single(index, chunk) for index, chunk in enumerate(chunks)]
        )
        return [item for item in ordered_results if item is not None]

    def _save_run_result(
        self,
        *,
        parsed_contract,
        chunk_results: list[dict[str, Any]],
        elapsed_seconds: float,
    ) -> str:
        """将 demo 运行结果写入 result 目录。"""
        result_dir = ensure_multi_agent_demo_result_dir(self.demo_config)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_stem = Path(parsed_contract.file_path).stem
        result_path = result_dir / f"linear_pipeline_{file_stem}_{timestamp}.txt"
        modification_count = sum(
            len(item["modifications"]) for item in chunk_results
        )

        lines: list[str] = [
            "线性顺序流水线 Demo 运行结果",
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"源文件路径：{parsed_contract.file_path}",
            f"源文件类型：{parsed_contract.file_type}",
            f"分块数量：{len(chunk_results)}",
            f"修改建议总数：{modification_count}",
            f"总计耗时（秒）：{elapsed_seconds:.3f}",
            "",
        ]

        for chunk_result in chunk_results:
            lines.append(f"## 分块{chunk_result['chunk_index']}")
            lines.append("### 分块内容")
            lines.append(chunk_result["chunk_text"])
            lines.append("")
            if not chunk_result["modifications"]:
                lines.append("### 审阅结果")
                lines.append("未发现修改建议")
                lines.append("")
                continue

            lines.append("### 审阅结果")
            for mod_index, modification in enumerate(chunk_result["modifications"], start=1):
                lines.append(f"- 修改点序号：{mod_index}")
                lines.append(f"  位置：{modification.get('position', '')}")
                lines.append(f"  原文：{modification.get('original_content', '')}")
                lines.append(f"  风险分析：{modification.get('risk_analysis', '')}")
                lines.append(f"  风险等级：{modification.get('risk_level', '')}")
                lines.append(f"  修改后的内容：{modification.get('suggested_content', '')}")
                lines.append(f"  修改理由：{modification.get('reason', '')}")
                lines.append(f"  风险类型：{modification.get('risk_type', '')}")
                lines.append("")

        with open(result_path, "w", encoding="utf-8") as file:
            file.write("\n".join(lines).strip() + "\n")
        return str(result_path)


class _FakeChatCompletions:
    """用于本地结构验证的假模型。"""

    def create(self, *, model: str, messages: list, tools: list, **kwargs):
        fake_review_result = """
【修改点1】付款条款表述不清
【原文】乙方完成工作后甲方付款。
【风险分析】付款触发条件和付款时间不明确，容易引发履约争议。
【风险等级】中
【修改后的内容】乙方完成工作并经甲方书面验收通过后，甲方应于10个工作日内支付合同款项。
【修改理由】补足验收条件和付款期限，减少争议。
【风险类型】付款条款
""".strip()
        return SimpleNamespace(
            choices=[
                SimpleNamespace(
                    message=SimpleNamespace(content=fake_review_result),
                )
            ]
        )


class _FakeLLM:
    """用于本地结构验证的假客户端。"""

    def __init__(self):
        self.chat = SimpleNamespace(completions=_FakeChatCompletions())


async def _main_test_linear_pipeline_demo():
    """本地最小结构验证。"""
    with tempfile.TemporaryDirectory(prefix="linear_pipeline_demo_test_") as temp_dir:
        sample_path = Path(temp_dir) / "sample_contract.docx"
        result_dir = Path(temp_dir) / "result"
        document = Document()
        document.add_paragraph("乙方完成工作后甲方付款。")
        document.add_paragraph("任何一方违约应承担责任。")
        document.save(sample_path)

        demo_config = MultiAgentDemoConfig(
            model=MultiAgentDemoModelConfig(
                model_name="fake-linear-pipeline-model",
                api_key="fake-key",
                api_base="https://example.com/v1",
                temperature=0.0,
                top_p=1.0,
                max_tokens=512,
            ),
            chunk_size=1000,
            max_concurrent_reviews=2,
            result_dir=str(result_dir),
        )
        service = LinearPipelineReviewDemo(
            config=demo_config,
            llm_client=_FakeLLM(),
        )

        result = await service.run(str(sample_path))
        print("linear_pipeline_demo self test result:")
        print(
            {
                "chunk_count": result["chunk_count"],
                "modification_count": result["modification_count"],
                "result_file_exists": Path(result["result_file_path"]).exists(),
            }
        )
        assert result["chunk_count"] == 1
        assert result["modification_count"] == 1
        assert Path(result["result_file_path"]).exists()
        print("linear_pipeline_demo self test passed")


if __name__ == "__main__":
    asyncio.run(_main_test_linear_pipeline_demo())
