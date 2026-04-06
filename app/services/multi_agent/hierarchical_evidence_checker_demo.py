"""层级主从 + 共享证据层 + 终审校验设计范式 Demo。

多 agent 链路说明：
1. 确定性预处理节点
   - 解析合同、按既有 toolkit 分块。
   - 为每个分块生成稳定 chunk_id。
   - 抽取轻量定义项、引用标记、主题标签，形成共享证据快照。
2. 主 agent（planner / supervisor）
   - 不直接浓缩全文语义。
   - 仅负责为每个分块选择合适 specialist，并生成任务列表。
   - 控制全局状态与执行顺序。
3. 依赖解析节点（deterministic）
   - 为每个 specialist task 组装 context package。
   - context package 包含：当前分块、相邻窗口、相关定义、相关分块片段、可用依赖句柄。
4. 并行 specialist agents
   - specialist 基于主 agent 指令和共享证据层并发审阅。
   - specialist 只输出结构化 findings，不直接改写全局状态。
5. 汇总层（merge）
   - 只合并结构化 findings，不重写底层证据。
6. checker agent
   - 校验高风险结论是否具备证据锚点。
   - 校验依赖句柄是否可解析。
   - 对不完整结论标记为“需复核”。

说明：
- 本文件为 demo，仅演示多 agent 编排结构，不接入真实 RAG。
- 共享证据层采用内存快照，提供类似检索句柄的体验。
- 真实模型运行时，最终结果会写入 app/services/multi_agent/result/。
"""

from __future__ import annotations

import argparse
import asyncio
import json
import re
import tempfile
import time
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from types import SimpleNamespace
from typing import Any

from docx import Document

try:
    from .config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        ensure_multi_agent_demo_result_dir,
        get_multi_agent_demo_config,
    )
    from .review_toolkit import MultiAgentReviewToolkit
except ImportError:
    from app.services.multi_agent.config import (
        MultiAgentDemoConfig,
        MultiAgentDemoModelConfig,
        ensure_multi_agent_demo_result_dir,
        get_multi_agent_demo_config,
    )
    from app.services.multi_agent.review_toolkit import MultiAgentReviewToolkit


RISK_SCORE_MAP = {"高": 3, "中": 2, "低": 1}
BLANK_PATTERN = re.compile(r"(留空|空白|待定|另行约定|____+|□|□\s*待填)")
REFERENCE_PATTERN = re.compile(r"(第[一二三四五六七八九十0-9]+条|附件[一二三四五六七八九十0-9A-Za-z]+|附表[一二三四五六七八九十0-9A-Za-z]+)")


@dataclass(slots=True)
class SpecialistProfile:
    """specialist 角色配置。"""

    specialist_name: str
    focus_area: str
    instruction: str
    keywords: list[str] = field(default_factory=list)


SPECIALIST_PROFILES = [
    SpecialistProfile(
        specialist_name="付款与结算 specialist",
        focus_area="付款与结算",
        instruction="重点审查付款触发条件、付款期限、验收联动、发票与结算条件。",
        keywords=["付款", "支付", "价款", "结算", "发票", "预付款", "进度款", "验收"],
    ),
    SpecialistProfile(
        specialist_name="责任与违约 specialist",
        focus_area="责任与违约",
        instruction="重点审查违约责任、赔偿范围、违约金、免责条款和责任分配。",
        keywords=["责任", "违约", "赔偿", "损失", "免责", "违约金", "索赔"],
    ),
    SpecialistProfile(
        specialist_name="终止与争议 specialist",
        focus_area="终止与争议",
        instruction="重点审查解除、终止、争议解决、通知、适用法律等条款。",
        keywords=["解除", "终止", "争议", "仲裁", "法院", "通知", "法律"],
    ),
    SpecialistProfile(
        specialist_name="定义与异常 specialist",
        focus_area="定义与异常",
        instruction="重点审查定义项引用、附件回指、空白字段、待定表述和格式性异常。",
        keywords=["定义", "系指", "是指", "指", "附件", "附表", "留空", "空白", "待定"],
    ),
]


@dataclass(slots=True)
class EvidenceNode:
    """共享证据层中的分块节点。"""

    chunk_id: str
    chunk_index: int
    text: str
    keyword_tags: list[str] = field(default_factory=list)
    reference_tokens: list[str] = field(default_factory=list)
    prev_chunk_id: str = ""
    next_chunk_id: str = ""


@dataclass(slots=True)
class DefinitionRecord:
    """轻量定义项记录。"""

    definition_id: str
    term: str
    description: str
    source_chunk_id: str


@dataclass(slots=True)
class EvidenceSnapshot:
    """共享证据快照。"""

    snapshot_id: str
    file_path: str
    file_type: str
    nodes: list[EvidenceNode] = field(default_factory=list)
    definitions: list[DefinitionRecord] = field(default_factory=list)
    available_handles: list[str] = field(default_factory=list)


@dataclass(slots=True)
class ContextSnippet:
    """context package 中的证据片段。"""

    handle_id: str
    label: str
    text: str


@dataclass(slots=True)
class ContextPackage:
    """specialist 可读的最小必要上下文包。"""

    chunk_id: str
    current_text: str
    adjacent_snippets: list[ContextSnippet] = field(default_factory=list)
    related_definition_snippets: list[ContextSnippet] = field(default_factory=list)
    related_evidence_snippets: list[ContextSnippet] = field(default_factory=list)
    dependency_ids: list[str] = field(default_factory=list)
    retrieval_handles: list[str] = field(default_factory=list)


@dataclass(slots=True)
class SpecialistTask:
    """planner 发给 specialist 的任务。"""

    task_id: str
    specialist_name: str
    focus_area: str
    instruction: str
    chunk_id: str
    context_package: ContextPackage


@dataclass(slots=True)
class SpecialistFinding:
    """specialist 输出的结构化发现。"""

    finding_id: str
    task_id: str
    specialist_name: str
    source_chunk_id: str
    title: str
    risk_level: str
    issue: str
    suggestion: str
    evidence: str
    evidence_chunk_ids: list[str] = field(default_factory=list)
    dependency_ids: list[str] = field(default_factory=list)
    tags: list[str] = field(default_factory=list)
    raw_output: str = ""


@dataclass(slots=True)
class MergedFinding:
    """merge 层归并后的发现。"""

    finding_key: str
    title: str
    risk_level: str
    issue: str
    suggestion: str
    evidence: str
    evidence_chunk_ids: list[str] = field(default_factory=list)
    dependency_ids: list[str] = field(default_factory=list)
    source_chunk_ids: list[str] = field(default_factory=list)
    supporting_specialists: list[str] = field(default_factory=list)
    tags: list[str] = field(default_factory=list)


@dataclass(slots=True)
class CheckedFinding:
    """checker 校验后的最终发现。"""

    finding_key: str
    title: str
    risk_level: str
    issue: str
    suggestion: str
    evidence: str
    evidence_chunk_ids: list[str] = field(default_factory=list)
    dependency_ids: list[str] = field(default_factory=list)
    source_chunk_ids: list[str] = field(default_factory=list)
    supporting_specialists: list[str] = field(default_factory=list)
    tags: list[str] = field(default_factory=list)
    checker_status: str = "accepted"
    checker_notes: list[str] = field(default_factory=list)


class HierarchicalEvidenceCheckerDemo(MultiAgentReviewToolkit):
    """层级主从 + 共享证据层 + checker 的多 agent demo。"""

    max_specialists_per_chunk = 2
    related_chunk_limit = 2

    async def run(
        self,
        file_path: str,
        *,
        stance: str | None = None,
        intensity: str | None = None,
        contract_type: str | None = None,
    ) -> dict[str, Any]:
        """执行该设计范式 demo。"""
        start_time = time.perf_counter()
        parsed_contract = await self.parse_contract_file(file_path)
        chunks = self.split_contract_text(parsed_contract.text)
        snapshot = self._build_evidence_snapshot(
            file_path=parsed_contract.file_path,
            file_type=parsed_contract.file_type,
            chunks=chunks,
        )
        planned_tasks = self._plan_specialist_tasks(snapshot)
        specialist_findings = await self._run_specialist_tasks(
            snapshot=snapshot,
            tasks=planned_tasks,
            stance=stance or self.demo_config.default_stance,
            intensity=intensity or self.demo_config.default_intensity,
            contract_type=contract_type or self.demo_config.default_contract_type,
        )
        merged_findings = self._merge_findings(specialist_findings)
        checked_findings = self._run_checker(snapshot, merged_findings)
        elapsed_seconds = time.perf_counter() - start_time
        result_paths = self._save_run_result(
            snapshot=snapshot,
            tasks=planned_tasks,
            specialist_findings=specialist_findings,
            merged_findings=merged_findings,
            checked_findings=checked_findings,
            elapsed_seconds=elapsed_seconds,
        )

        accepted_count = sum(
            1 for item in checked_findings if item.checker_status == "accepted"
        )
        recheck_count = sum(
            1 for item in checked_findings if item.checker_status == "needs_recheck"
        )

        return {
            "demo_type": "hierarchical_evidence_checker",
            "file_path": snapshot.file_path,
            "file_type": snapshot.file_type,
            "snapshot_id": snapshot.snapshot_id,
            "chunk_count": len(snapshot.nodes),
            "definition_count": len(snapshot.definitions),
            "task_count": len(planned_tasks),
            "specialist_finding_count": len(specialist_findings),
            "merged_finding_count": len(merged_findings),
            "final_finding_count": len(checked_findings),
            "accepted_finding_count": accepted_count,
            "needs_recheck_count": recheck_count,
            "elapsed_seconds": elapsed_seconds,
            "result_file_path": result_paths["text"],
            "result_json_path": result_paths["json"],
        }

    def _build_evidence_snapshot(
        self,
        *,
        file_path: str,
        file_type: str,
        chunks: list[str],
    ) -> EvidenceSnapshot:
        """构建 demo 用共享证据快照。"""
        nodes: list[EvidenceNode] = []
        definitions: list[DefinitionRecord] = []

        for index, chunk_text in enumerate(chunks, start=1):
            chunk_id = f"chunk-{index}"
            keyword_tags = self._infer_chunk_tags(chunk_text)
            reference_tokens = self._extract_reference_tokens(chunk_text)
            prev_chunk_id = f"chunk-{index - 1}" if index > 1 else ""
            next_chunk_id = f"chunk-{index + 1}" if index < len(chunks) else ""
            nodes.append(
                EvidenceNode(
                    chunk_id=chunk_id,
                    chunk_index=index,
                    text=chunk_text,
                    keyword_tags=keyword_tags,
                    reference_tokens=reference_tokens,
                    prev_chunk_id=prev_chunk_id,
                    next_chunk_id=next_chunk_id,
                )
            )
            for term, description in self._extract_definition_pairs(chunk_text):
                definitions.append(
                    DefinitionRecord(
                        definition_id=f"definition-{len(definitions) + 1}",
                        term=term,
                        description=description,
                        source_chunk_id=chunk_id,
                    )
                )

        handles = [node.chunk_id for node in nodes]
        handles.extend(item.definition_id for item in definitions)
        handles.extend(
            token
            for node in nodes
            for token in node.reference_tokens
        )
        return EvidenceSnapshot(
            snapshot_id=f"snapshot-{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            file_path=file_path,
            file_type=file_type,
            nodes=nodes,
            definitions=definitions,
            available_handles=self._unique_texts(handles),
        )

    def _plan_specialist_tasks(
        self,
        snapshot: EvidenceSnapshot,
    ) -> list[SpecialistTask]:
        """planner 根据分块内容为 specialist 派单。"""
        tasks: list[SpecialistTask] = []
        for node in snapshot.nodes:
            scored_profiles: list[tuple[int, SpecialistProfile]] = []
            for profile in SPECIALIST_PROFILES:
                score = self._score_profile(node.text, profile.keywords)
                if score > 0:
                    scored_profiles.append((score, profile))

            scored_profiles.sort(key=lambda item: (-item[0], item[1].specialist_name))
            selected_profiles = [
                item[1] for item in scored_profiles[: self.max_specialists_per_chunk]
            ]

            if not selected_profiles:
                selected_profiles = [
                    SpecialistProfile(
                        specialist_name="通用风险 specialist",
                        focus_area="通用风险",
                        instruction="补充识别当前分块中的通用合同风险与明显遗漏。",
                        keywords=[],
                    )
                ]

            if BLANK_PATTERN.search(node.text) and not any(
                item.specialist_name == "定义与异常 specialist"
                for item in selected_profiles
            ):
                selected_profiles.append(SPECIALIST_PROFILES[3])

            for profile in selected_profiles:
                task_id = f"{node.chunk_id}-{self._normalize_key(profile.focus_area)}"
                context_package = self._build_context_package(
                    snapshot=snapshot,
                    node=node,
                    profile=profile,
                )
                tasks.append(
                    SpecialistTask(
                        task_id=task_id,
                        specialist_name=profile.specialist_name,
                        focus_area=profile.focus_area,
                        instruction=profile.instruction,
                        chunk_id=node.chunk_id,
                        context_package=context_package,
                    )
                )
        return tasks

    def _build_context_package(
        self,
        *,
        snapshot: EvidenceSnapshot,
        node: EvidenceNode,
        profile: SpecialistProfile,
    ) -> ContextPackage:
        """依赖解析节点为 specialist 组装上下文包。"""
        node_map = {item.chunk_id: item for item in snapshot.nodes}
        adjacent_snippets: list[ContextSnippet] = []
        dependency_ids: list[str] = []

        for chunk_id, label in (
            (node.prev_chunk_id, "上文窗口"),
            (node.next_chunk_id, "下文窗口"),
        ):
            if not chunk_id or chunk_id not in node_map:
                continue
            target_node = node_map[chunk_id]
            adjacent_snippets.append(
                ContextSnippet(
                    handle_id=target_node.chunk_id,
                    label=label,
                    text=self._truncate_text(target_node.text),
                )
            )
            dependency_ids.append(target_node.chunk_id)

        related_definition_snippets: list[ContextSnippet] = []
        for definition in snapshot.definitions:
            if definition.term and definition.term in node.text:
                related_definition_snippets.append(
                    ContextSnippet(
                        handle_id=definition.definition_id,
                        label=f"定义项：{definition.term}",
                        text=self._truncate_text(definition.description),
                    )
                )
                dependency_ids.extend(
                    [definition.definition_id, definition.source_chunk_id]
                )

        related_evidence_snippets: list[ContextSnippet] = []
        profile_keywords = set(profile.keywords)
        candidate_nodes: list[tuple[int, EvidenceNode]] = []
        for candidate in snapshot.nodes:
            if candidate.chunk_id == node.chunk_id:
                continue
            overlap = profile_keywords.intersection(candidate.keyword_tags)
            if overlap:
                candidate_nodes.append((len(overlap), candidate))

        candidate_nodes.sort(key=lambda item: (-item[0], item[1].chunk_index))
        for _, candidate in candidate_nodes[: self.related_chunk_limit]:
            related_evidence_snippets.append(
                ContextSnippet(
                    handle_id=candidate.chunk_id,
                    label="相关分块",
                    text=self._truncate_text(candidate.text),
                )
            )
            dependency_ids.append(candidate.chunk_id)

        for token in node.reference_tokens:
            dependency_ids.append(token)

        retrieval_handles = [node.chunk_id]
        retrieval_handles.extend(item.handle_id for item in adjacent_snippets)
        retrieval_handles.extend(
            item.handle_id for item in related_definition_snippets
        )
        retrieval_handles.extend(item.handle_id for item in related_evidence_snippets)
        retrieval_handles.extend(node.reference_tokens)

        return ContextPackage(
            chunk_id=node.chunk_id,
            current_text=node.text,
            adjacent_snippets=adjacent_snippets,
            related_definition_snippets=related_definition_snippets,
            related_evidence_snippets=related_evidence_snippets,
            dependency_ids=self._unique_texts(dependency_ids),
            retrieval_handles=self._unique_texts(retrieval_handles),
        )

    async def _run_specialist_tasks(
        self,
        *,
        snapshot: EvidenceSnapshot,
        tasks: list[SpecialistTask],
        stance: str,
        intensity: str,
        contract_type: str,
    ) -> list[SpecialistFinding]:
        """并发执行 specialist 审阅任务。"""
        semaphore = asyncio.Semaphore(self.demo_config.max_concurrent_reviews)
        ordered_results: list[list[SpecialistFinding] | None] = [None] * len(tasks)

        async def _run_single(index: int, task: SpecialistTask):
            async with semaphore:
                ordered_results[index] = await self._run_specialist_task(
                    snapshot=snapshot,
                    task=task,
                    stance=stance,
                    intensity=intensity,
                    contract_type=contract_type,
                )

        await asyncio.gather(
            *[_run_single(index, task) for index, task in enumerate(tasks)]
        )
        flattened: list[SpecialistFinding] = []
        for item in ordered_results:
            if item:
                flattened.extend(item)
        return flattened

    async def _run_specialist_task(
        self,
        *,
        snapshot: EvidenceSnapshot,
        task: SpecialistTask,
        stance: str,
        intensity: str,
        contract_type: str,
    ) -> list[SpecialistFinding]:
        """单个 specialist 的执行逻辑。"""
        prompt = self._build_specialist_prompt(
            snapshot=snapshot,
            task=task,
            stance=stance,
            intensity=intensity,
            contract_type=contract_type,
        )
        messages = [
            {
                "role": "system",
                "content": "你是合同审阅多 agent 系统中的 specialist agent，只能基于给定上下文包输出纯 JSON。",
            },
            {"role": "user", "content": prompt},
        ]
        raw_output = ""
        try:
            response = await asyncio.to_thread(
                self.llm.chat.completions.create,
                model=self.model_config.model_name,
                messages=messages,
                tools=[],
                temperature=self.model_config.temperature,
                top_p=self.model_config.top_p,
                max_tokens=self.model_config.max_tokens,
            )
            raw_output = response.choices[0].message.content or ""
            return self._parse_specialist_output(
                task=task,
                raw_output=raw_output,
            )
        except Exception:
            return await self._fallback_specialist_output(
                task=task,
                raw_output=raw_output,
                stance=stance,
                intensity=intensity,
                contract_type=contract_type,
            )

    def _build_specialist_prompt(
        self,
        *,
        snapshot: EvidenceSnapshot,
        task: SpecialistTask,
        stance: str,
        intensity: str,
        contract_type: str,
    ) -> str:
        """构造 specialist prompt。"""
        context_package = task.context_package
        adjacent_text = self._render_context_snippets(context_package.adjacent_snippets)
        definition_text = self._render_context_snippets(
            context_package.related_definition_snippets
        )
        related_text = self._render_context_snippets(
            context_package.related_evidence_snippets
        )
        handles = "、".join(context_package.retrieval_handles) or "无"
        dependencies = "、".join(context_package.dependency_ids) or "无"

        return (
            "请严格输出 JSON，不要输出 markdown，不要输出解释文字。\n"
            "输出结构如下：\n"
            "{\n"
            '  "task_summary": "字符串",\n'
            '  "findings": [\n'
            "    {\n"
            '      "title": "问题标题",\n'
            '      "risk_level": "高/中/低",\n'
            '      "issue": "问题说明",\n'
            '      "suggestion": "修改建议",\n'
            '      "evidence": "关键证据原文",\n'
            '      "evidence_chunk_ids": ["chunk-1"],\n'
            '      "dependency_ids": ["chunk-2", "definition-1"],\n'
            '      "tags": ["标签1", "标签2"]\n'
            "    }\n"
            "  ]\n"
            "}\n"
            "要求：\n"
            "1. 只能依据当前任务提供的上下文包下结论，不得编造证据。\n"
            "2. 若没有明确问题，findings 返回空数组。\n"
            "3. 高风险结论必须尽量补齐 evidence 与 evidence_chunk_ids。\n"
            "4. dependency_ids 只能填写当前可用依赖句柄中的值。\n"
            f"5. 当前快照版本：{snapshot.snapshot_id}。\n"
            f"6. 用户立场：{stance}；审阅强度：{intensity}；合同类型：{contract_type}。\n"
            f"任务编号：{task.task_id}\n"
            f"specialist：{task.specialist_name}\n"
            f"审阅焦点：{task.focus_area}\n"
            f"执行指令：{task.instruction}\n"
            f"当前分块标识：{context_package.chunk_id}\n"
            "当前分块原文：\n"
            f"{context_package.current_text}\n"
            "相邻窗口：\n"
            f"{adjacent_text}\n"
            "相关定义：\n"
            f"{definition_text}\n"
            "相关证据片段：\n"
            f"{related_text}\n"
            f"可用依赖句柄：{handles}\n"
            f"建议优先核验的依赖：{dependencies}\n"
        )

    def _parse_specialist_output(
        self,
        *,
        task: SpecialistTask,
        raw_output: str,
    ) -> list[SpecialistFinding]:
        """解析 specialist JSON 输出。"""
        payload = self._extract_json_payload(raw_output)
        raw_findings = payload.get("findings") or []
        results: list[SpecialistFinding] = []

        for index, item in enumerate(raw_findings, start=1):
            title = str(item.get("title", "")).strip() or f"{task.focus_area}问题{index}"
            risk_level = self._normalize_risk_level(item.get("risk_level"))
            issue = str(item.get("issue", "")).strip()
            suggestion = str(item.get("suggestion", "")).strip()
            evidence = str(item.get("evidence", "")).strip()
            evidence_chunk_ids = self._normalize_handle_list(
                item.get("evidence_chunk_ids"),
                prefix="chunk-",
            )
            dependency_ids = self._normalize_handle_list(
                item.get("dependency_ids"),
                prefix="",
            )
            tags = self._normalize_text_list(item.get("tags"))
            results.append(
                SpecialistFinding(
                    finding_id=f"{task.task_id}-finding-{index}",
                    task_id=task.task_id,
                    specialist_name=task.specialist_name,
                    source_chunk_id=task.chunk_id,
                    title=title,
                    risk_level=risk_level,
                    issue=issue,
                    suggestion=suggestion,
                    evidence=evidence,
                    evidence_chunk_ids=evidence_chunk_ids or [task.chunk_id],
                    dependency_ids=dependency_ids,
                    tags=tags,
                    raw_output=raw_output,
                )
            )
        return results

    async def _fallback_specialist_output(
        self,
        *,
        task: SpecialistTask,
        raw_output: str,
        stance: str,
        intensity: str,
        contract_type: str,
    ) -> list[SpecialistFinding]:
        """specialist 输出异常时退化为 toolkit 的普通审阅。"""
        fallback_findings: list[SpecialistFinding] = []
        parsed_items = self._parse_review_result(raw_output) if raw_output else []

        if not parsed_items:
            parsed_items = await self.review_chunk(
                chunk_text=task.context_package.current_text,
                stance=stance,
                intensity=intensity,
                contract_type=contract_type,
                context=f"{task.specialist_name} 对 {task.chunk_id} 的兜底审阅。",
            )

        for index, item in enumerate(parsed_items, start=1):
            fallback_findings.append(
                SpecialistFinding(
                    finding_id=f"{task.task_id}-fallback-{index}",
                    task_id=task.task_id,
                    specialist_name=task.specialist_name,
                    source_chunk_id=task.chunk_id,
                    title=str(item.get("risk_type", "")).strip() or f"{task.focus_area}问题{index}",
                    risk_level=self._normalize_risk_level(item.get("risk_level")),
                    issue=str(item.get("risk_analysis", "")).strip(),
                    suggestion=str(item.get("suggested_content", "")).strip(),
                    evidence=str(item.get("original_content", "")).strip(),
                    evidence_chunk_ids=[task.chunk_id],
                    dependency_ids=task.context_package.dependency_ids[:2],
                    tags=[task.focus_area],
                    raw_output=raw_output,
                )
            )
        return fallback_findings

    def _merge_findings(
        self,
        findings: list[SpecialistFinding],
    ) -> list[MergedFinding]:
        """merge 层只做结构化归并。"""
        grouped: dict[str, list[SpecialistFinding]] = defaultdict(list)
        for item in findings:
            merge_key = self._normalize_key(item.title or item.issue[:24])
            grouped[merge_key].append(item)

        merged_results: list[MergedFinding] = []
        for merge_key, items in grouped.items():
            primary = max(
                items,
                key=lambda item: (
                    RISK_SCORE_MAP.get(item.risk_level, 0),
                    len(item.evidence),
                ),
            )
            merged_results.append(
                MergedFinding(
                    finding_key=merge_key,
                    title=primary.title,
                    risk_level=primary.risk_level,
                    issue="；".join(
                        self._unique_texts([item.issue for item in items if item.issue])
                    ),
                    suggestion="；".join(
                        self._unique_texts(
                            [item.suggestion for item in items if item.suggestion]
                        )
                    ),
                    evidence="；".join(
                        self._unique_texts([item.evidence for item in items if item.evidence])
                    ),
                    evidence_chunk_ids=self._unique_texts(
                        [
                            chunk_id
                            for item in items
                            for chunk_id in item.evidence_chunk_ids
                            if chunk_id
                        ]
                    ),
                    dependency_ids=self._unique_texts(
                        [
                            dependency_id
                            for item in items
                            for dependency_id in item.dependency_ids
                            if dependency_id
                        ]
                    ),
                    source_chunk_ids=self._unique_texts(
                        [item.source_chunk_id for item in items if item.source_chunk_id]
                    ),
                    supporting_specialists=self._unique_texts(
                        [item.specialist_name for item in items if item.specialist_name]
                    ),
                    tags=self._unique_texts(
                        [tag for item in items for tag in item.tags if tag]
                    ),
                )
            )

        merged_results.sort(
            key=lambda item: (
                -RISK_SCORE_MAP.get(item.risk_level, 0),
                item.title,
            )
        )
        return merged_results

    def _run_checker(
        self,
        snapshot: EvidenceSnapshot,
        merged_findings: list[MergedFinding],
    ) -> list[CheckedFinding]:
        """checker 对归并后的发现做最终校验。"""
        valid_chunk_ids = {node.chunk_id for node in snapshot.nodes}
        valid_dependency_ids = set(snapshot.available_handles).union(valid_chunk_ids)
        checked_results: list[CheckedFinding] = []

        for item in merged_findings:
            notes: list[str] = []
            invalid_evidence_ids = [
                chunk_id
                for chunk_id in item.evidence_chunk_ids
                if chunk_id not in valid_chunk_ids
            ]
            invalid_dependency_ids = [
                handle_id
                for handle_id in item.dependency_ids
                if handle_id not in valid_dependency_ids
            ]

            if item.risk_level == "高" and not item.evidence.strip():
                notes.append("高风险结论缺少明确证据。")
            if item.risk_level == "高" and not item.evidence_chunk_ids:
                notes.append("高风险结论缺少证据锚点。")
            if invalid_evidence_ids:
                notes.append(
                    f"证据锚点不可解析：{'、'.join(invalid_evidence_ids)}。"
                )
            if invalid_dependency_ids:
                notes.append(
                    f"依赖句柄不可解析：{'、'.join(invalid_dependency_ids)}。"
                )

            checker_status = "accepted" if not notes else "needs_recheck"
            checked_results.append(
                CheckedFinding(
                    finding_key=item.finding_key,
                    title=item.title,
                    risk_level=item.risk_level,
                    issue=item.issue,
                    suggestion=item.suggestion,
                    evidence=item.evidence,
                    evidence_chunk_ids=item.evidence_chunk_ids,
                    dependency_ids=item.dependency_ids,
                    source_chunk_ids=item.source_chunk_ids,
                    supporting_specialists=item.supporting_specialists,
                    tags=item.tags,
                    checker_status=checker_status,
                    checker_notes=notes,
                )
            )
        return checked_results

    def _save_run_result(
        self,
        *,
        snapshot: EvidenceSnapshot,
        tasks: list[SpecialistTask],
        specialist_findings: list[SpecialistFinding],
        merged_findings: list[MergedFinding],
        checked_findings: list[CheckedFinding],
        elapsed_seconds: float,
    ) -> dict[str, str]:
        """写入 text/json 结果文件。"""
        result_dir = ensure_multi_agent_demo_result_dir(self.demo_config)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_stem = Path(snapshot.file_path).stem
        text_path = (
            result_dir
            / f"hierarchical_evidence_checker_{file_stem}_{timestamp}.txt"
        )
        json_path = (
            result_dir
            / f"hierarchical_evidence_checker_{file_stem}_{timestamp}.json"
        )

        task_counter = Counter(task.specialist_name for task in tasks)
        accepted_count = sum(
            1 for item in checked_findings if item.checker_status == "accepted"
        )
        recheck_count = len(checked_findings) - accepted_count

        lines = [
            "层级主从 + 共享证据层 + 终审校验 Demo 运行结果",
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"源文件路径：{snapshot.file_path}",
            f"源文件类型：{snapshot.file_type}",
            f"证据快照版本：{snapshot.snapshot_id}",
            f"分块数量：{len(snapshot.nodes)}",
            f"定义项数量：{len(snapshot.definitions)}",
            f"任务数量：{len(tasks)}",
            f"specialist 原始发现数量：{len(specialist_findings)}",
            f"merge 后发现数量：{len(merged_findings)}",
            f"checker 最终发现数量：{len(checked_findings)}",
            f"checker 接受数量：{accepted_count}",
            f"checker 需复核数量：{recheck_count}",
            f"总计耗时（秒）：{elapsed_seconds:.3f}",
            "",
            "## specialist 任务分布",
        ]

        for specialist_name, count in sorted(task_counter.items()):
            lines.append(f"- {specialist_name}：{count}")

        lines.append("")
        lines.append("## checker 最终结论")
        if not checked_findings:
            lines.append("未形成最终发现。")
            lines.append("")
        else:
            for item in checked_findings:
                lines.append(f"### {item.title}")
                lines.append(f"状态：{item.checker_status}")
                lines.append(f"风险等级：{item.risk_level}")
                lines.append(f"来源分块：{','.join(item.source_chunk_ids) or '无'}")
                lines.append(
                    f"支持 specialist：{','.join(item.supporting_specialists) or '无'}"
                )
                lines.append(
                    f"证据锚点：{','.join(item.evidence_chunk_ids) or '无'}"
                )
                lines.append(
                    f"依赖句柄：{','.join(item.dependency_ids) or '无'}"
                )
                lines.append(f"问题说明：{item.issue or '无'}")
                lines.append(f"修改建议：{item.suggestion or '无'}")
                lines.append(f"证据摘要：{self._truncate_text(item.evidence, limit=180) or '无'}")
                if item.checker_notes:
                    lines.append(f"checker 说明：{'；'.join(item.checker_notes)}")
                lines.append("")

        with open(text_path, "w", encoding="utf-8") as file:
            file.write("\n".join(lines).strip() + "\n")

        json_payload = {
            "snapshot": {
                "snapshot_id": snapshot.snapshot_id,
                "file_path": snapshot.file_path,
                "file_type": snapshot.file_type,
                "chunk_count": len(snapshot.nodes),
                "definition_count": len(snapshot.definitions),
                "available_handles": snapshot.available_handles,
            },
            "tasks": [asdict(item) for item in tasks],
            "specialist_findings": [asdict(item) for item in specialist_findings],
            "merged_findings": [asdict(item) for item in merged_findings],
            "checked_findings": [asdict(item) for item in checked_findings],
            "elapsed_seconds": elapsed_seconds,
        }
        with open(json_path, "w", encoding="utf-8") as file:
            json.dump(json_payload, file, ensure_ascii=False, indent=2)

        return {"text": str(text_path), "json": str(json_path)}

    def _infer_chunk_tags(self, text: str) -> list[str]:
        """根据关键词为分块打轻量标签。"""
        tags: list[str] = []
        for profile in SPECIALIST_PROFILES:
            if any(keyword in text for keyword in profile.keywords):
                tags.extend(keyword for keyword in profile.keywords if keyword in text)
        if BLANK_PATTERN.search(text):
            tags.append("留空")
        return self._unique_texts(tags)

    def _extract_reference_tokens(self, text: str) -> list[str]:
        """抽取条款/附件引用标记。"""
        return self._unique_texts(REFERENCE_PATTERN.findall(text))

    def _extract_definition_pairs(self, text: str) -> list[tuple[str, str]]:
        """轻量抽取定义项。"""
        pairs: list[tuple[str, str]] = []
        for match in re.finditer(
            r"(?:“(?P<quoted>[^”]{2,24})”|(?P<plain>[\u4e00-\u9fa5A-Za-z]{2,24}))\s*(?:系指|是指|指)\s*(?P<desc>[^。；\n]{4,120})",
            text,
        ):
            term = (match.group("quoted") or match.group("plain") or "").strip()
            description = (match.group("desc") or "").strip()
            if term and description:
                pairs.append((term, description))
        return pairs

    def _score_profile(self, text: str, keywords: list[str]) -> int:
        """计算分块与 specialist 的匹配分数。"""
        return sum(1 for keyword in keywords if keyword in text)

    def _extract_json_payload(self, raw_output: str) -> dict[str, Any]:
        """从模型输出中提取 JSON 对象。"""
        text = raw_output.strip()
        if not text:
            raise ValueError("空输出")
        fenced_match = re.search(r"```json\s*(\{[\s\S]*\})\s*```", text)
        if fenced_match:
            text = fenced_match.group(1)
        else:
            object_match = re.search(r"\{[\s\S]*\}", text)
            if object_match:
                text = object_match.group(0)
        payload = json.loads(text)
        if not isinstance(payload, dict):
            raise ValueError("输出不是 JSON 对象")
        return payload

    def _normalize_risk_level(self, value: Any) -> str:
        """标准化风险等级。"""
        text = str(value).strip()
        if text in {"高", "中", "低"}:
            return text
        return "中"

    def _normalize_text_list(self, value: Any) -> list[str]:
        """清洗普通文本列表。"""
        if not isinstance(value, list):
            return []
        return self._unique_texts([str(item).strip() for item in value if str(item).strip()])

    def _normalize_handle_list(self, value: Any, *, prefix: str) -> list[str]:
        """清洗句柄列表。"""
        result: list[str] = []
        if not isinstance(value, list):
            return result
        for item in value:
            text = str(item).strip()
            if not text:
                continue
            if prefix and not text.startswith(prefix):
                continue
            result.append(text)
        return self._unique_texts(result)

    def _normalize_key(self, value: str) -> str:
        """标准化分组键。"""
        return re.sub(r"\s+", "", str(value).strip()).lower()

    def _render_context_snippets(self, snippets: list[ContextSnippet]) -> str:
        """渲染上下文片段。"""
        if not snippets:
            return "无"
        lines: list[str] = []
        for snippet in snippets:
            lines.append(
                f"- {snippet.handle_id}（{snippet.label}）：{snippet.text}"
            )
        return "\n".join(lines)

    def _truncate_text(self, text: str, *, limit: int = 120) -> str:
        """截断长文本，避免在结果中写入过长原文。"""
        clean_text = re.sub(r"\s+", " ", text.strip())
        if len(clean_text) <= limit:
            return clean_text
        return clean_text[: limit - 3] + "..."

    def _unique_texts(self, items: list[str]) -> list[str]:
        """保持顺序去重。"""
        result: list[str] = []
        for item in items:
            text = str(item).strip()
            if text and text not in result:
                result.append(text)
        return result


class _FakeChatCompletions:
    """用于本地结构验证的假模型。"""

    def create(self, *, model: str, messages: list, tools: list, **kwargs):
        user_prompt = messages[-1]["content"]

        if (
            "付款与结算 specialist" in user_prompt
            and "30个工作日内支付全部合同价款" in user_prompt
        ):
            content = json.dumps(
                {
                    "task_summary": "付款条款依赖验收定义，需要补充付款时限合理性审查。",
                    "findings": [
                        {
                            "title": "付款期限偏长且依赖验收定义",
                            "risk_level": "高",
                            "issue": "付款义务以验收合格为前提，但付款期限达到30个工作日，可能拉长回款周期。",
                            "suggestion": "建议缩短至10至15个工作日，并明确书面验收确认流程。",
                            "evidence": "甲方应在验收合格后30个工作日内支付全部合同价款。",
                            "evidence_chunk_ids": ["chunk-2"],
                            "dependency_ids": ["chunk-1", "definition-1"],
                            "tags": ["付款", "验收"],
                        }
                    ],
                },
                ensure_ascii=False,
            )
        elif (
            "责任与违约 specialist" in user_prompt
            and "违约金比例留空" in user_prompt
        ):
            content = json.dumps(
                {
                    "task_summary": "违约责任条款存在空白项。",
                    "findings": [
                        {
                            "title": "违约金比例留空",
                            "risk_level": "高",
                            "issue": "违约责任条款提到违约金比例，但关键比例未填写。",
                            "suggestion": "建议明确违约金比例、计算基数和触发条件。",
                            "evidence": "违约责任：任何一方违约均应承担损失，违约金比例留空。",
                            "evidence_chunk_ids": ["chunk-3"],
                            "dependency_ids": ["chunk-3"],
                            "tags": ["违约", "留空"],
                        }
                    ],
                },
                ensure_ascii=False,
            )
        elif (
            "终止与争议 specialist" in user_prompt
            and "项目所在地法院" in user_prompt
        ):
            content = json.dumps(
                {
                    "task_summary": "争议解决条款需要核验表述是否完整。",
                    "findings": [
                        {
                            "title": "争议解决地需结合项目管辖核验",
                            "risk_level": "中",
                            "issue": "条款直接指定项目所在地法院，需结合实际管辖规则与双方便利性评估。",
                            "suggestion": "建议补充更明确的争议解决层级，并确认法院管辖约定的有效性。",
                            "evidence": "争议提交项目所在地法院。",
                            "evidence_chunk_ids": ["chunk-3"],
                            "dependency_ids": ["chunk-3"],
                            "tags": ["争议解决"],
                        }
                    ],
                },
                ensure_ascii=False,
            )
        else:
            content = json.dumps(
                {"task_summary": "当前任务未识别出明确问题。", "findings": []},
                ensure_ascii=False,
            )

        return SimpleNamespace(
            choices=[
                SimpleNamespace(
                    message=SimpleNamespace(content=content),
                )
            ]
        )


class _FakeLLM:
    """用于本地结构验证的假客户端。"""

    def __init__(self):
        self.chat = SimpleNamespace(completions=_FakeChatCompletions())


async def _main_test_hierarchical_evidence_checker_demo():
    """本地最小结构验证。"""
    with tempfile.TemporaryDirectory(
        prefix="hierarchical_evidence_checker_demo_test_"
    ) as temp_dir:
        sample_path = Path(temp_dir) / "sample_contract.docx"
        result_dir = Path(temp_dir) / "result"
        document = Document()
        document.add_paragraph(
            "定义条款：本合同所称“验收合格”是指甲方书面确认工作成果符合约定标准。"
        )
        document.add_paragraph(
            "付款条款：甲方应在验收合格后30个工作日内支付全部合同价款。"
        )
        document.add_paragraph(
            "违约责任：任何一方违约均应承担损失，违约金比例留空。争议提交项目所在地法院。"
        )
        document.save(sample_path)

        demo_config = MultiAgentDemoConfig(
            model=MultiAgentDemoModelConfig(
                model_name="fake-hierarchical-evidence-model",
                api_key="fake-key",
                api_base="https://example.com/v1",
                temperature=0.0,
                top_p=1.0,
                max_tokens=1024,
            ),
            chunk_size=40,
            max_concurrent_reviews=2,
            result_dir=str(result_dir),
        )
        service = HierarchicalEvidenceCheckerDemo(
            config=demo_config,
            llm_client=_FakeLLM(),
        )

        result = await service.run(str(sample_path))
        print("hierarchical_evidence_checker_demo self test result:")
        print(
            {
                "chunk_count": result["chunk_count"],
                "task_count": result["task_count"],
                "final_finding_count": result["final_finding_count"],
                "accepted_finding_count": result["accepted_finding_count"],
                "elapsed_seconds": round(result["elapsed_seconds"], 3),
                "result_file_exists": Path(result["result_file_path"]).exists(),
                "result_json_exists": Path(result["result_json_path"]).exists(),
            }
        )
        assert result["chunk_count"] >= 3
        assert result["task_count"] >= 3
        assert result["final_finding_count"] >= 2
        assert result["accepted_finding_count"] >= 1
        assert result["elapsed_seconds"] >= 0
        assert Path(result["result_file_path"]).exists()
        assert Path(result["result_json_path"]).exists()
        print("hierarchical_evidence_checker_demo self test passed")


async def _main_run_real_model_demo(file_path: str):
    """使用真实模型运行 demo。"""
    config = get_multi_agent_demo_config()
    service = HierarchicalEvidenceCheckerDemo(config=config)
    result = await service.run(file_path)
    print("hierarchical_evidence_checker_demo real run result:")
    print(
        {
            "chunk_count": result["chunk_count"],
            "task_count": result["task_count"],
            "final_finding_count": result["final_finding_count"],
            "accepted_finding_count": result["accepted_finding_count"],
            "needs_recheck_count": result["needs_recheck_count"],
            "elapsed_seconds": round(result["elapsed_seconds"], 3),
            "result_file_path": result["result_file_path"],
            "result_json_path": result["result_json_path"],
        }
    )


def _build_arg_parser() -> argparse.ArgumentParser:
    """构造命令行参数。"""
    parser = argparse.ArgumentParser(
        description="层级主从 + 共享证据层 + 终审校验 Demo"
    )
    parser.add_argument(
        "--file",
        dest="file_path",
        help="真实模型运行时的合同文件路径（doc/docx）",
    )
    return parser


if __name__ == "__main__":
    arguments = _build_arg_parser().parse_args()
    if arguments.file_path:
        asyncio.run(_main_run_real_model_demo(arguments.file_path))
    else:
        asyncio.run(_main_test_hierarchical_evidence_checker_demo())
