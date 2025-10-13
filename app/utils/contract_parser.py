from typing import Dict
from langchain.schema import SystemMessage, HumanMessage
from docx import Document
from typing import List
import pdfplumber
from pdf2image import convert_from_path
import pytesseract
import json
import os

def read_docx_text(file_path: str) -> str:
    """安全读取 docx 文件的文本内容，支持段落和表格"""
    try:
        doc = Document(file_path)
        texts: List[str] = []

        # 读取段落
        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text.strip())

        # 读取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        texts.append(cell_text)

        return "\n".join(texts)

    except Exception as e:
        raise RuntimeError(f"❌ 读取 Word 文件失败: {e}")


def read_pdf_first_page(file_path: str) -> str:
    """读取 PDF 文件第一页内容，支持文字 PDF 和扫描件 PDF"""
    try:
        # 1. 尝试用 pdfplumber 提取文本
        text = ""
        with pdfplumber.open(file_path) as pdf:
            if len(pdf.pages) > 0:
                text = pdf.pages[0].extract_text() or ""

        if text.strip():
            return text.strip()

        # 2. 如果提取为空，说明可能是扫描件，使用 OCR
        pages = convert_from_path(file_path, dpi=300, first_page=1, last_page=1)
        if not pages:
            return ""

        ocr_text = pytesseract.image_to_string(pages[0], lang="chi_sim+eng")
        return ocr_text.strip()

    except Exception as e:
        print(f"读取 PDF 失败: {e}")
        return ""

async def extract_parties_with_llm(file_path: str, llm) -> Dict[str, str]:
    """使用 LLM 提取合同中的甲方和乙方"""
    try:
        ext = os.path.splitext(file_path)[1].lower()

        # 1.根据文件类型读取
        if ext == ".pdf":
            text = read_pdf_first_page(file_path)
        else:
            text = read_docx_text(file_path)
        sample_text = text[:200]  # 取前200字分析

        # print(sample_text)

        # 2. 构造提示词
        base_prompt = f"""
你是一名专业的合同律师。请从以下合同内容中提取出甲方和乙方的名称（如果有简称，也一并标注）。
请严格以 JSON 格式输出结果，不要包含其他文字、说明或 Markdown 代码块。输出示例如下：
{{
  "party_a": "xxx公司",
  "party_b": "yyy公司"
}}
如果未找到，则用空字符串代替。
"""
        # 3. 构造消息列表（与审阅部分一致）
        messages = [
            SystemMessage(content="你是一个专业的合同信息提取助手。"),
            HumanMessage(content=f"{base_prompt}\n\n合同内容:\n{sample_text}")
        ]

        # 4. 调用大模型
        response = llm.invoke(messages)
        raw_output = response.content.strip()

        # 5. 尝试解析模型输出
        try:
            result = json.loads(raw_output)
        except json.JSONDecodeError:
            # 如果模型输出不是严格的JSON，可以做一次兜底处理
            result = {"party_a": "", "party_b": ""}
        return {
            "party_a": result.get("party_a", ""),
            "party_b": result.get("party_b", "")
        }

    except Exception as e:
        print(f"提取甲乙方失败: {e}")
        return {"party_a": "", "party_b": ""}
