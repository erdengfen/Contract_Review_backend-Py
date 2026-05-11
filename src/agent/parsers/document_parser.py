"""文档解析路由入口。

本文件负责把 Step 4 的 `DocumentSensingResult` 接到实际解析器，不负责切片或审阅。
"""
from __future__ import annotations

import tempfile
from pathlib import Path

from src.agent.contracts.document import DocumentSensingResult, ParsedDocument
from src.agent.parsers.doc_normalizer import DocNormalizationError, normalize_doc_to_docx
from src.agent.parsers.docx_parser import parse_docx_file
from src.agent.parsers.pdf_parser import MinerULightParser, parse_text_pdf_file


def parse_document_from_sensing(
    sensing: DocumentSensingResult,
    *,
    mineru_client: MinerULightParser | None = None,
) -> ParsedDocument:
    """根据文档感知结果调用对应解析器。"""

    if sensing.trust.blocked or sensing.route == "blocked":
        return _failed_from_sensing(sensing, ["document_sensing_blocked"])
    if sensing.route == "parse_docx_pending":
        return _parse_docx_from_sensing(sensing)
    if sensing.route == "format_normalization_required":
        return _parse_doc_after_normalization(sensing)
    if sensing.route == "pdf_text_extract":
        return _parse_text_pdf_from_sensing(sensing, mineru_client)
    return _failed_from_sensing(sensing, [f"parser_route_not_supported:{sensing.route}"])


def _parse_docx_from_sensing(sensing: DocumentSensingResult) -> ParsedDocument:
    """解析已通过感知层分流的 DOCX 文件。"""

    parsed = parse_docx_file(
        sensing.file_reference.save_path,
        filename=sensing.file_reference.original_filename,
        file_id=sensing.file_reference.file_id,
        original_file_path=sensing.file_reference.save_path,
    )
    parsed.warnings = [*sensing.warnings, *parsed.warnings]
    parsed.metadata["source_route"] = sensing.route
    return parsed


def _parse_doc_after_normalization(sensing: DocumentSensingResult) -> ParsedDocument:
    """执行 DOC 归一化后复用 DOCX 结构解析器。"""

    try:
        with tempfile.TemporaryDirectory(prefix="agent_doc_normalize_") as temp_dir:
            docx_path = normalize_doc_to_docx(sensing.file_reference.save_path, temp_dir)
            parsed = parse_docx_file(
                str(docx_path),
                filename=sensing.file_reference.original_filename,
                file_id=sensing.file_reference.file_id,
                original_file_path=sensing.file_reference.save_path,
            )
    except DocNormalizationError as error:
        return _failed_from_sensing(sensing, ["doc_normalization_failed", str(error)])
    parsed.file_type = "doc"
    parsed.normalized_file_path = None
    parsed.warnings = [*sensing.warnings, "doc_normalized_to_docx_ephemeral", *parsed.warnings]
    parsed.metadata["source_route"] = sensing.route
    parsed.metadata["normalization"] = "doc_to_docx"
    return parsed


def _parse_text_pdf_from_sensing(
    sensing: DocumentSensingResult,
    mineru_client: MinerULightParser | None,
) -> ParsedDocument:
    """解析普通 PDF，保留 MinerU 轻量解析可选接入。"""

    parsed = parse_text_pdf_file(
        sensing.file_reference.save_path,
        filename=sensing.file_reference.original_filename,
        file_id=sensing.file_reference.file_id,
        original_file_path=sensing.file_reference.save_path,
        mineru_client=mineru_client,
    )
    parsed.warnings = [*sensing.warnings, *parsed.warnings]
    parsed.metadata["source_route"] = sensing.route
    return parsed


def _failed_from_sensing(sensing: DocumentSensingResult, warnings: list[str]) -> ParsedDocument:
    """根据感知结果构造解析失败文档。"""

    return ParsedDocument(
        filename=sensing.file_reference.original_filename,
        file_type=sensing.detected_file_type,
        file_path=sensing.file_reference.save_path,
        contract_content=None,
        blocks=[],
        parser_name="agent-parser-router",
        parse_status="failed",
        warnings=[*sensing.warnings, *warnings],
        metadata={
            "source_route": sensing.route,
            "content_format": "markdown",
            "block_count": 0,
        },
    )


def _main_test_document_parser() -> None:
    """执行解析路由入口的本文件自检。"""

    from docx import Document

    from src.agent.contracts.document import DocumentIntakeRequest
    from src.agent.document_intake.sensing import sense_document

    with tempfile.TemporaryDirectory(prefix="agent_parser_route_") as temp_dir:
        path = Path(temp_dir) / "sample.docx"
        document = Document()
        document.add_paragraph("合同正文")
        document.save(path)
        request = DocumentIntakeRequest(filename="sample.docx", file_type="docx", save_path=str(path))
        parsed = parse_document_from_sensing(sense_document(request))
        assert parsed.parse_status == "success"
        assert parsed.blocks[0].text == "合同正文"


if __name__ == "__main__":
    _main_test_document_parser()
