"""DOCX 结构解析器。

本文件按 DOCX 原始 body 顺序解析段落、标题、列表、表格和图片占位，输出解析层结构。
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocumentFactory
from docx.oxml import CT_P, CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from src.agent.contracts.document import DocumentBlock, ParsedDocument, SourceLocation


@dataclass
class _ParseState:
    """保存单个 DOCX 解析过程中的顺序、页码和文本偏移。"""

    file_id: int | None
    file_path: str
    page_number: int | None
    page_number_source: str
    block_index: int = 0
    cursor: int = 0
    section_title: str | None = None


def parse_docx_file(
    file_path: str,
    *,
    filename: str | None = None,
    file_id: int | None = None,
    original_file_path: str | None = None,
    normalized_file_path: str | None = None,
) -> ParsedDocument:
    """解析 DOCX 文件，返回结构化文档和 Markdown 视图。"""

    path = Path(file_path).expanduser().resolve()
    source_path = original_file_path or str(path)
    if not path.exists():
        return _failed_document(filename or path.name, "docx", source_path, ["docx_file_missing"])

    document = DocxDocumentFactory(str(path))
    has_page_breaks = _document_has_explicit_page_breaks(document)
    state = _build_parse_state(file_id, source_path, has_page_breaks)
    blocks, content_parts = _parse_body_blocks(document, state)
    warnings = _build_parse_warnings(blocks, has_page_breaks)
    return ParsedDocument(
        filename=filename or path.name,
        file_type="docx",
        file_path=source_path,
        normalized_file_path=normalized_file_path,
        contract_content="\n\n".join(content_parts) if content_parts else None,
        blocks=blocks,
        parser_name="python-docx-structure",
        parse_status="success" if blocks else "failed",
        warnings=warnings,
        metadata=_build_document_metadata(blocks, state.page_number_source),
    )


def _failed_document(filename: str, file_type: str, file_path: str, warnings: list[str]) -> ParsedDocument:
    """构造解析失败文档。"""

    return ParsedDocument(
        filename=filename,
        file_type=file_type,
        file_path=file_path,
        parser_name="python-docx-structure",
        parse_status="failed",
        warnings=warnings,
        metadata={"content_format": "markdown", "block_count": 0},
    )


def _build_parse_state(file_id: int | None, file_path: str, has_page_breaks: bool) -> _ParseState:
    """根据是否存在显式分页符创建解析状态。"""

    return _ParseState(
        file_id=file_id,
        file_path=file_path,
        page_number=1 if has_page_breaks else None,
        page_number_source="explicit_page_break" if has_page_breaks else "unavailable",
    )


def _parse_body_blocks(document: Any, state: _ParseState) -> tuple[list[DocumentBlock], list[str]]:
    """按 DOCX body 原始顺序解析内容块。"""

    blocks: list[DocumentBlock] = []
    content_parts: list[str] = []
    for block_type, body_block in _iter_body_blocks(document):
        parsed_block = _parse_body_block(block_type, body_block, state)
        if parsed_block:
            blocks.append(parsed_block)
            content_parts.append(parsed_block.text)
        if block_type == "paragraph" and _paragraph_has_page_break(body_block):
            _advance_page_after_break(state)
    return blocks, content_parts


def _iter_body_blocks(document: Any) -> list[tuple[str, Paragraph | Table]]:
    """遍历 DOCX body，保留段落和表格在原文中的顺序。"""

    blocks: list[tuple[str, Paragraph | Table]] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(("paragraph", Paragraph(child, document)))
        elif isinstance(child, CT_Tbl):
            blocks.append(("table", Table(child, document)))
    return blocks


def _parse_body_block(
    block_type: str,
    body_block: Paragraph | Table,
    state: _ParseState,
) -> DocumentBlock | None:
    """解析单个 body 元素。"""

    if block_type == "paragraph":
        return _parse_paragraph_block(body_block, state)
    return _parse_table_block(body_block, state)


def _parse_paragraph_block(paragraph: Paragraph, state: _ParseState) -> DocumentBlock | None:
    """解析段落、标题、列表和图片占位。"""

    text = paragraph.text.strip()
    images = _extract_image_placeholders(paragraph)
    if not text and not images:
        return None
    element_type = _paragraph_element_type(paragraph, images)
    if element_type == "heading" and text:
        state.section_title = text
    markdown = _paragraph_to_markdown(text, element_type, images)
    metadata = _base_metadata(element_type, state.page_number_source)
    if images:
        metadata["images"] = images
    return _make_block(state, markdown, text or markdown, element_type, metadata)


def _parse_table_block(table: Table, state: _ParseState) -> DocumentBlock | None:
    """解析表格为 Markdown 视图和结构化单元格。"""

    rows = _extract_table_rows(table)
    markdown = _table_rows_to_markdown(rows)
    if not markdown:
        return None
    metadata = _base_metadata("table", state.page_number_source)
    metadata["table"] = {"rows": rows}
    return _make_block(state, markdown, _table_rows_to_plain_text(rows), "table", metadata)


def _make_block(
    state: _ParseState,
    text: str,
    normalized_text: str,
    element_type: str,
    metadata: dict[str, Any],
) -> DocumentBlock:
    """创建带源文档定位的解析块。"""

    state.block_index += 1
    start_offset = state.cursor
    end_offset = start_offset + len(text)
    state.cursor = end_offset + 2
    location = SourceLocation(
        file_path=state.file_path,
        page_number=state.page_number,
        block_index=state.block_index,
        layout_index=state.block_index,
        start_offset=start_offset,
        end_offset=end_offset,
        element_type=element_type,
    )
    return DocumentBlock(
        block_id=_build_block_id(state),
        text=text,
        normalized_text=normalized_text,
        source_location=location,
        section_title=state.section_title,
        page_number=state.page_number,
        block_index=state.block_index,
        char_count=len(text),
        metadata=metadata,
    )


def _build_block_id(state: _ParseState) -> str:
    """生成解析块唯一标识。"""

    file_part = f"file-{state.file_id}" if state.file_id is not None else Path(state.file_path).stem
    page_part = state.page_number if state.page_number is not None else "unknown"
    return f"{file_part}-page-{page_part}-block-{state.block_index}"


def _paragraph_element_type(paragraph: Paragraph, images: list[dict[str, str]]) -> str:
    """判断段落元素类型。"""

    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if style_name.startswith("heading") or style_name.startswith("标题"):
        return "heading"
    if _paragraph_is_list_item(paragraph):
        return "list_item"
    if images and not paragraph.text.strip():
        return "image"
    return "paragraph"


def _paragraph_to_markdown(text: str, element_type: str, images: list[dict[str, str]]) -> str:
    """把段落元素转为面向切片的 Markdown 视图。"""

    parts: list[str] = []
    if text:
        parts.append(_format_paragraph_text(text, element_type))
    for image in images:
        parts.append(f"[图片占位:{image['relationship_id']}]")
    return "\n".join(parts)


def _format_paragraph_text(text: str, element_type: str) -> str:
    """按元素类型格式化段落文本。"""

    if element_type == "heading":
        return f"# {text}"
    if element_type == "list_item":
        return f"- {text}"
    return text


def _paragraph_is_list_item(paragraph: Paragraph) -> bool:
    """判断段落是否包含 Word 列表编号信息。"""

    return "w:numPr" in paragraph._element.xml


def _extract_image_placeholders(paragraph: Paragraph) -> list[dict[str, str]]:
    """提取段落内图片关系，生成不可执行占位元数据。"""

    images: list[dict[str, str]] = []
    for run in paragraph.runs:
        for drawing in run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
            image = _image_metadata_from_drawing(paragraph, drawing)
            if image:
                images.append(image)
    return images


def _image_metadata_from_drawing(paragraph: Paragraph, drawing: Any) -> dict[str, str] | None:
    """从 drawing 节点提取图片关系信息。"""

    blip = drawing.find(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
    if blip is None:
        return None
    relationship_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    if not relationship_id:
        return None
    relation = paragraph.part.rels.get(relationship_id)
    content_type = getattr(getattr(relation, "target_part", None), "content_type", "unknown")
    return {"relationship_id": relationship_id, "content_type": content_type}


def _extract_table_rows(table: Table) -> list[list[dict[str, Any]]]:
    """提取表格单元格文本和合并线索。"""

    rows: list[list[dict[str, Any]]] = []
    for row_index, row in enumerate(table.rows):
        cells = []
        for column_index, cell in enumerate(row.cells):
            merge_info = _cell_merge_info(cell)
            cells.append(
                {
                    "row_index": row_index,
                    "column_index": column_index,
                    "text": _clean_cell_text(cell.text),
                    **merge_info,
                }
            )
        rows.append(cells)
    return rows


def _cell_merge_info(cell: _Cell) -> dict[str, Any]:
    """提取单元格横向和纵向合并信息。"""

    table_cell_properties = cell._tc.tcPr
    if table_cell_properties is None:
        return {"grid_span": 1, "vertical_merge": None}
    grid_span = getattr(getattr(table_cell_properties, "gridSpan", None), "val", 1)
    vertical_merge = getattr(getattr(table_cell_properties, "vMerge", None), "val", None)
    return {"grid_span": int(grid_span or 1), "vertical_merge": vertical_merge}


def _clean_cell_text(text: str) -> str:
    """清理单元格文本，避免破坏 Markdown 表格。"""

    return " ".join(text.replace("|", "\\|").split())


def _table_rows_to_markdown(rows: list[list[dict[str, Any]]]) -> str:
    """把结构化表格行转换为 Markdown 表格。"""

    if not rows:
        return ""
    width = max((len(row) for row in rows), default=0)
    if width == 0:
        return ""
    normalized_rows = [_normalize_table_row(row, width) for row in rows]
    header = normalized_rows[0]
    separator = ["---"] * width
    lines = [_markdown_row(header), _markdown_row(separator)]
    for row in normalized_rows[1:]:
        lines.append(_markdown_row(row))
    return "\n".join(lines)


def _normalize_table_row(row: list[dict[str, Any]], width: int) -> list[str]:
    """将表格行补齐到统一列数。"""

    values = [cell["text"] for cell in row]
    return (values + [""] * width)[:width]


def _markdown_row(values: list[str]) -> str:
    """生成单行 Markdown 表格。"""

    return "| " + " | ".join(values) + " |"


def _table_rows_to_plain_text(rows: list[list[dict[str, Any]]]) -> str:
    """生成表格的纯文本视图。"""

    return "\n".join("\t".join(cell["text"] for cell in row) for row in rows)


def _base_metadata(element_type: str, page_number_source: str) -> dict[str, Any]:
    """生成所有解析块共享的元数据。"""

    return {
        "element_type": element_type,
        "page_number_source": page_number_source,
        "content_format": "markdown",
    }


def _document_has_explicit_page_breaks(document: Any) -> bool:
    """判断 DOCX 是否包含显式分页符。"""

    return any(_paragraph_has_page_break(paragraph) for paragraph in document.paragraphs)


def _paragraph_has_page_break(paragraph: Paragraph) -> bool:
    """判断段落是否包含显式分页符或渲染分页标记。"""

    xml = paragraph._element.xml
    return 'w:type="page"' in xml or "lastRenderedPageBreak" in xml


def _advance_page_after_break(state: _ParseState) -> None:
    """遇到显式分页符后推进后续解析块页码。"""

    if state.page_number is None:
        state.page_number = 2
        return
    state.page_number += 1


def _build_parse_warnings(blocks: list[DocumentBlock], has_page_breaks: bool) -> list[str]:
    """生成非敏感解析告警。"""

    warnings: list[str] = []
    if not blocks:
        warnings.append("docx_no_text_block")
    if not has_page_breaks:
        warnings.append("docx_page_number_unavailable")
    return warnings


def _build_document_metadata(blocks: list[DocumentBlock], page_number_source: str) -> dict[str, Any]:
    """生成文档级结构化元数据。"""

    return {
        "content_format": "markdown",
        "block_count": len(blocks),
        "page_number_source": page_number_source,
    }


def _main_test_docx_parser() -> None:
    """执行 DOCX 解析器的本文件自检。"""

    import tempfile

    with tempfile.TemporaryDirectory(prefix="agent_docx_parser_") as temp_dir:
        path = Path(temp_dir) / "sample.docx"
        document = DocxDocumentFactory()
        document.add_paragraph("第一条 总则", style="Heading 1")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "节点"
        table.cell(0, 1).text = "金额"
        table.cell(1, 0).text = "验收后"
        table.cell(1, 1).text = "100"
        document.save(path)
        parsed = parse_docx_file(str(path), file_id=1)
        assert parsed.parse_status == "success"
        assert parsed.blocks[0].metadata["element_type"] == "heading"
        assert parsed.blocks[1].metadata["element_type"] == "table"


if __name__ == "__main__":
    _main_test_docx_parser()
