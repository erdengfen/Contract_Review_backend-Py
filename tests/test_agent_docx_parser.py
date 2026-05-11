"""Agent DOCX 解析层测试。

本测试验证 Step 5.1 DOCX 结构解析器和 Step 4 感知结果到解析层的接线。
"""
import json
from pathlib import Path

import pytest
from docx import Document

from src.agent.contracts import DocumentIntakeRequest
from src.agent.document_intake import sense_document
from src.agent.parsers import parse_document_from_sensing, parse_docx_file


# 真实 DOCX 验证入口：把文件放到这里，或直接修改这个常量指向你的样本。
TEST_DATA_DIR = Path(__file__).resolve().parent / "data"
REAL_DOCX_PATH = TEST_DATA_DIR / "real_contract.docx"
PARSED_OUTPUT_DIR = TEST_DATA_DIR / "parsed_outputs"


def _write_structured_docx(path: Path) -> None:
    """写入包含标题、表格和正文的 DOCX 样例。"""

    document = Document()
    document.add_paragraph("第一条 总则", style="Heading 1")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "付款节点"
    table.cell(0, 1).text = "金额"
    table.cell(1, 0).text = "验收后"
    table.cell(1, 1).text = "100万元"
    document.add_paragraph("表后条款继续说明。")
    document.save(path)


def _write_page_break_docx(path: Path) -> None:
    """写入包含显式分页符的 DOCX 样例。"""

    document = Document()
    document.add_paragraph("第一页条款")
    document.add_page_break()
    document.add_paragraph("第二页条款")
    document.save(path)


def _build_request(path: Path, filename: str, file_type: str) -> DocumentIntakeRequest:
    """构造文档感知请求。"""

    return DocumentIntakeRequest(
        file_id=88,
        user_id=7,
        filename=filename,
        file_type=file_type,
        save_path=str(path),
    )


def _dump_parsed_document_for_review(parsed, source_path: Path) -> tuple[Path, Path]:
    """把真实样本解析结果临时落盘，便于人工审查 Markdown 和结构化 JSON。"""

    PARSED_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    md_path = PARSED_OUTPUT_DIR / f"{source_path.stem}.parsed.md"
    json_path = PARSED_OUTPUT_DIR / f"{source_path.stem}.parsed.json"
    md_path.write_text(parsed.contract_content or "", encoding="utf-8")
    json_path.write_text(
        json.dumps(parsed.model_dump(mode="json"), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return md_path, json_path


def test_docx_parser_preserves_paragraph_table_order(tmp_path: Path) -> None:
    """DOCX 解析器应按原始 body 顺序输出标题、表格和正文。"""

    path = tmp_path / "合同.docx"
    _write_structured_docx(path)
    parsed = parse_docx_file(str(path), filename="合同.docx", file_id=88)

    assert parsed.parse_status == "success"
    assert [block.metadata["element_type"] for block in parsed.blocks] == ["heading", "table", "paragraph"]
    assert parsed.contract_content.index("# 第一条 总则") < parsed.contract_content.index("| 付款节点 | 金额 |")
    assert parsed.contract_content.index("| 验收后 | 100万元 |") < parsed.contract_content.index("表后条款")


def test_docx_parser_outputs_structured_table_metadata(tmp_path: Path) -> None:
    """表格块应同时包含 Markdown 文本和结构化单元格视图。"""

    path = tmp_path / "合同.docx"
    _write_structured_docx(path)
    parsed = parse_docx_file(str(path), filename="合同.docx", file_id=88)
    table_block = parsed.blocks[1]

    assert table_block.text.startswith("| 付款节点 | 金额 |")
    assert table_block.metadata["table"]["rows"][1][0]["text"] == "验收后"
    assert table_block.source_location.layout_index == 2
    assert table_block.source_location.element_type == "table"


def test_docx_parser_marks_explicit_page_break_source(tmp_path: Path) -> None:
    """显式分页符存在时，解析层应记录页码和页码来源。"""

    path = tmp_path / "分页合同.docx"
    _write_page_break_docx(path)
    parsed = parse_docx_file(str(path), filename="分页合同.docx", file_id=88)

    assert parsed.blocks[0].page_number == 1
    assert parsed.blocks[1].page_number == 2
    assert parsed.metadata["page_number_source"] == "explicit_page_break"
    assert parsed.blocks[0].metadata["page_number_source"] == "explicit_page_break"


def test_step4_docx_route_connects_to_parser(tmp_path: Path) -> None:
    """Step 4 的 DOCX 路由应能直接接到 Step 5.1 解析层。"""

    path = tmp_path / "合同.docx"
    _write_structured_docx(path)
    sensing = sense_document(_build_request(path, "合同.docx", "docx"))
    parsed = parse_document_from_sensing(sensing)

    assert sensing.route == "parse_docx_pending"
    assert parsed.parse_status == "success"
    assert parsed.metadata["source_route"] == "parse_docx_pending"
    assert parsed.blocks[0].block_id == "file-88-page-unknown-block-1"


def test_doc_route_is_wired_to_normalization_failure_without_app_dependency(tmp_path: Path) -> None:
    """DOC 路由应进入归一化入口，失败时返回解析层失败结构。"""

    path = tmp_path / "旧合同.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0" + b"\x00" * 32)
    sensing = sense_document(_build_request(path, "旧合同.doc", "doc"))
    parsed = parse_document_from_sensing(sensing)

    assert sensing.route == "format_normalization_required"
    assert parsed.parse_status == "failed"
    assert "doc_normalization_failed" in parsed.warnings
    assert parsed.metadata["source_route"] == "format_normalization_required"


def test_real_docx_file_can_be_parsed_and_dumped_for_review() -> None:
    """解析 tests/data 下的真实 DOCX，并把审查产物写到 tests/data/parsed_outputs。"""

    if not REAL_DOCX_PATH.exists():
        pytest.skip(f"请将真实 DOCX 文件放到: {REAL_DOCX_PATH}")

    sensing = sense_document(_build_request(REAL_DOCX_PATH, REAL_DOCX_PATH.name, "docx"))
    parsed = parse_document_from_sensing(sensing)
    md_path, json_path = _dump_parsed_document_for_review(parsed, REAL_DOCX_PATH)

    assert sensing.route == "parse_docx_pending"
    assert parsed.parse_status == "success"
    assert parsed.blocks
    assert md_path.exists()
    assert json_path.exists()


if __name__ == "__main__":
    raise SystemExit(pytest.main([__file__]))
